#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MEJORADOR DE ESTILOS COMPLETO APA 7
Versión avanzada para mejorar estilos de todo el documento

Basado en el editor_documento_apa7.py pero enfocado en presentación visual
"""

from editor_documento_apa7 import EditorDocumentoAPA7
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json
from datetime import datetime

class MejoradorEstilosCompleto(EditorDocumentoAPA7):
    """Mejorador especializado en estilos y presentación visual"""
    
    def __init__(self, ruta_documento_original):
        super().__init__(ruta_documento_original)
        self.secciones_procesadas = []
        
    def mejorar_estilos_documento_completo(self):
        """Mejora estilos de todo el documento por secciones"""
        print("\n🎨 INICIANDO MEJORA COMPLETA DE ESTILOS DEL DOCUMENTO...")
        
        # Abrir documento
        if self.ruta_copia.exists():
            self.documento = Document(str(self.ruta_copia))
        else:
            self.fase1_preparacion_y_analisis()
        
        # Identificar todas las secciones principales
        secciones = self._identificar_secciones_principales()
        
        resumen_mejoras = {
            "secciones_procesadas": [],
            "total_mejoras": {
                "titulos_formateados": 0,
                "subtitulos_mejorados": 0,
                "parrafos_ajustados": 0,
                "listas_formateadas": 0
            },
            "timestamp": datetime.now().isoformat()
        }
        
        # Procesar cada sección
        for seccion in secciones:
            print(f"\n📚 Procesando sección: {seccion['nombre']}")
            mejoras_seccion = self._mejorar_estilos_seccion(
                seccion['inicio'], 
                seccion['fin'], 
                seccion['nombre']
            )
            
            seccion['mejoras'] = mejoras_seccion
            resumen_mejoras["secciones_procesadas"].append(seccion)
            
            # Acumular totales
            resumen_mejoras["total_mejoras"]["titulos_formateados"] += len(mejoras_seccion.get("titulos_principales_formateados", []))
            resumen_mejoras["total_mejoras"]["subtitulos_mejorados"] += len(mejoras_seccion.get("subtitulos_mejorados", []))
            resumen_mejoras["total_mejoras"]["parrafos_ajustados"] += len(mejoras_seccion.get("parrafos_ajustados", []))
            resumen_mejoras["total_mejoras"]["listas_formateadas"] += len(mejoras_seccion.get("listas_formateadas", []))
        
        # Guardar documento
        self.documento.save(str(self.ruta_copia))
        
        print("\n✅ MEJORAS COMPLETAS APLICADAS A TODO EL DOCUMENTO")
        return resumen_mejoras
    
    def _identificar_secciones_principales(self):
        """Identifica todas las secciones principales del documento"""
        secciones = []
        seccion_actual = None
        
        patrones_secciones = {
            'INTRODUCCIÓN': ['INTRODUCCIÓN', 'INTRODUCTION', 'I. INTRODUCCIÓN'],
            'OBJETIVOS': ['OBJETIVOS', 'OBJETIVO', 'II. OBJETIVOS'],
            'METODOLOGÍA': ['METODOLOGÍA', 'METODOLOGIA', 'III. METODOLOGÍA'],
            'RESULTADOS': ['RESULTADOS', 'HALLAZGOS', 'IV. RESULTADOS'],
            'IMPLICACIONES': ['IMPLICACIONES', 'APLICACIONES', 'V. IMPLICACIONES'],
            'CONCLUSIONES': ['CONCLUSIONES', 'CONCLUSIÓN', 'VI. CONCLUSIONES'],
            'REFERENCIAS': ['REFERENCIAS', 'BIBLIOGRAFÍA', 'REFERENCIAS BIBLIOGRÁFICAS']
        }
        
        for i, paragraph in enumerate(self.documento.paragraphs):
            texto = paragraph.text.strip().upper()
            
            for nombre_seccion, patrones in patrones_secciones.items():
                if any(patron in texto for patron in patrones) and len(texto) < 200:
                    # Finalizar sección anterior si existe
                    if seccion_actual:
                        seccion_actual['fin'] = i - 1
                        secciones.append(seccion_actual)
                    
                    # Iniciar nueva sección
                    seccion_actual = {
                        'nombre': nombre_seccion,
                        'inicio': i,
                        'fin': None,
                        'titulo_encontrado': texto
                    }
                    break
        
        # Finalizar última sección
        if seccion_actual:
            seccion_actual['fin'] = len(self.documento.paragraphs) - 1
            secciones.append(seccion_actual)
        
        print(f"📋 Identificadas {len(secciones)} secciones principales:")
        for seccion in secciones:
            print(f"  📍 {seccion['nombre']}: párrafos {seccion['inicio']}-{seccion['fin']}")
        
        return secciones
    
    def _mejorar_estilos_seccion(self, inicio, fin, nombre_seccion):
        """Mejora estilos de una sección específica"""
        
        cambios_estilos = {
            "nombre_seccion": nombre_seccion,
            "rango": f"{inicio}-{fin}",
            "titulos_principales_formateados": [],
            "subtitulos_mejorados": [],
            "parrafos_ajustados": [],
            "listas_formateadas": []
        }
        
        for i in range(inicio, fin + 1):
            if i >= len(self.documento.paragraphs):
                break
                
            paragraph = self.documento.paragraphs[i]
            texto_original = paragraph.text.strip()
            
            if not texto_original:
                continue
            
            # 1. TÍTULOS PRINCIPALES
            if self._es_titulo_principal_general(texto_original, nombre_seccion):
                self._aplicar_estilo_titulo_principal(paragraph)
                cambios_estilos["titulos_principales_formateados"].append({
                    "posicion": i,
                    "texto": texto_original[:50] + "..." if len(texto_original) > 50 else texto_original
                })
            
            # 2. SUBTÍTULOS
            elif self._es_subtitulo_general(texto_original):
                nivel = self._determinar_nivel_subtitulo_general(texto_original)
                self._aplicar_estilo_subtitulo(paragraph, nivel)
                cambios_estilos["subtitulos_mejorados"].append({
                    "posicion": i,
                    "texto": texto_original[:40] + "..." if len(texto_original) > 40 else texto_original,
                    "nivel": nivel
                })
            
            # 3. PÁRRAFOS DE CONTENIDO
            elif len(texto_original) > 20:
                self._mejorar_formato_parrafo_contenido(paragraph)
                cambios_estilos["parrafos_ajustados"].append({
                    "posicion": i,
                    "mejoras": "Formato APA 7 aplicado"
                })
            
            # 4. LISTAS
            if self._contiene_lista(texto_original):
                self._formatear_lista_apa7(paragraph)
                cambios_estilos["listas_formateadas"].append({
                    "posicion": i,
                    "tipo": "Lista APA 7"
                })
        
        return cambios_estilos
    
    def _es_titulo_principal_general(self, texto, seccion):
        """Detecta títulos principales de cualquier sección"""
        texto_upper = texto.upper()
        
        # Patrones generales de títulos principales
        patrones_titulo = [
            'INTRODUCCIÓN', 'OBJETIVOS', 'METODOLOGÍA', 'RESULTADOS', 
            'IMPLICACIONES', 'CONCLUSIONES', 'REFERENCIAS'
        ]
        
        return (any(patron in texto_upper for patron in patrones_titulo) and 
                len(texto) < 100 and 
                not any(char in texto for char in ['\t', '.', ':', ';']))
    
    def _es_subtitulo_general(self, texto):
        """Detecta subtítulos en cualquier sección"""
        texto_upper = texto.upper()
        
        # Patrones de subtítulos
        indicadores = [
            # Numeración
            texto.startswith(('1.', '2.', '3.', '4.', '5.', '6.')),
            texto.startswith(('1.1', '1.2', '2.1', '2.2', '3.1', '3.2')),
            texto.startswith(('4.1', '4.2', '4.3', '4.4', '5.1', '5.2')),
            
            # Palabras clave de subtítulos
            any(palabra in texto_upper for palabra in [
                'ENFOQUE', 'DISEÑO', 'ANÁLISIS', 'CARACTERIZACIÓN',
                'POBLACIÓN', 'MUESTRA', 'TÉCNICAS', 'HERRAMIENTAS',
                'LIMITACIONES', 'VALIDACIÓN', 'TIPOLOGÍA', 'PERFIL'
            ]),
            
            # Formato típico de subtítulos
            len(texto) < 150 and texto.endswith(':') == False and '\t' not in texto
        ]
        
        return any(indicadores) and len(texto) > 10
    
    def _determinar_nivel_subtitulo_general(self, texto):
        """Determina nivel jerárquico de subtítulos"""
        
        # Nivel 2: Subtítulos principales con numeración simple
        if any(texto.startswith(num) for num in ['1.', '2.', '3.', '4.', '5.', '6.']):
            return 2
        
        # Nivel 3: Sub-subtítulos con numeración decimal
        elif any(texto.startswith(num) for num in ['1.1', '1.2', '2.1', '2.2', '3.1', '3.2', '4.1', '4.2']):
            return 3
        
        # Nivel 2 por defecto para subtítulos sin numeración clara
        return 2

def demo_mejora_completa():
    """Demostración de mejora completa del documento"""
    
    ruta_documento = r"assets\Consultoria\01 Informe en elaboración\01 Informe Tipologías IIEE 2025 1.0.docx"
    
    try:
        print("🎨 DEMO: MEJORA COMPLETA DE ESTILOS DEL DOCUMENTO")
        print("="*70)
        
        # Crear mejorador
        mejorador = MejoradorEstilosCompleto(ruta_documento)
        
        # Aplicar mejoras completas
        resumen = mejorador.mejorar_estilos_documento_completo()
        
        # Mostrar resumen
        print("\n" + "="*70)
        print("📊 RESUMEN FINAL DE MEJORAS APLICADAS:")
        print("="*70)
        print(f"🎯 Total títulos formateados: {resumen['total_mejoras']['titulos_formateados']}")
        print(f"📚 Total subtítulos mejorados: {resumen['total_mejoras']['subtitulos_mejorados']}")
        print(f"📝 Total párrafos ajustados: {resumen['total_mejoras']['parrafos_ajustados']}")
        print(f"📋 Total listas formateadas: {resumen['total_mejoras']['listas_formateadas']}")
        
        print(f"\n📚 Secciones procesadas: {len(resumen['secciones_procesadas'])}")
        for seccion in resumen['secciones_procesadas']:
            print(f"  📍 {seccion['nombre']}: {len(seccion['mejoras']['parrafos_ajustados'])} párrafos mejorados")
        
        # Guardar reporte completo
        archivo_reporte = Path("mejoras_estilos_completo.json")
        with open(archivo_reporte, 'w', encoding='utf-8') as f:
            json.dump(resumen, f, indent=2, ensure_ascii=False)
        
        print(f"\n💾 Documento con estilos mejorados: {mejorador.ruta_copia}")
        print(f"📊 Reporte completo: {archivo_reporte}")
        print("\n🎉 ¡MEJORA COMPLETA FINALIZADA!")
        
        return mejorador
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

if __name__ == "__main__":
    demo_mejora_completa()
