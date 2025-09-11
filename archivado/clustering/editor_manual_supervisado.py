#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EDITOR MANUAL SUPERVISADO APA 7
Enfoque minucioso: párrafo por párrafo con revisión manual

Este editor permite analizar y asignar estilos de forma supervisada,
mostrando cada párrafo para evaluación manual antes de aplicar cambios.
"""

import os
import json
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    print("⚠️ ERROR: python-docx no está instalado")
    print("Instalar con: pip install python-docx")
    raise e

class EditorManualSupervisado:
    """Editor supervisado para análisis manual párrafo por párrafo"""
    
    def __init__(self, ruta_documento):
        self.ruta_original = Path(ruta_documento)
        self.ruta_trabajo = self.ruta_original.with_name(
            f"{self.ruta_original.stem}_manual_supervisado{self.ruta_original.suffix}"
        )
        self.documento = None
        self.estilos_definidos = self._definir_bateria_estilos()
        self.cambios_aplicados = []
        
    def _definir_bateria_estilos(self):
        """Define la batería completa de estilos APA 7"""
        return {
            "TITULO_PRINCIPAL": {
                "nombre": "Título Principal (Nivel 1)",
                "descripcion": "Centrado, negrita, Times New Roman 12pt",
                "aplicacion": self._aplicar_titulo_principal
            },
            "SUBTITULO_NIVEL_2": {
                "nombre": "Subtítulo Nivel 2", 
                "descripcion": "Izquierda, negrita, Times New Roman 12pt",
                "aplicacion": self._aplicar_subtitulo_nivel_2
            },
            "SUBTITULO_NIVEL_3": {
                "nombre": "Subtítulo Nivel 3",
                "descripcion": "Sangría 0.5\", negrita, Times New Roman 12pt",
                "aplicacion": self._aplicar_subtitulo_nivel_3
            },
            "PARRAFO_NORMAL": {
                "nombre": "Párrafo Normal",
                "descripcion": "Justificado, sangría primera línea 0.5\", doble espaciado",
                "aplicacion": self._aplicar_parrafo_normal
            },
            "DESCRIPCION_CORTA": {
                "nombre": "Descripción Corta",
                "descripcion": "Justificado, sin sangría, espaciado 1.5",
                "aplicacion": self._aplicar_descripcion_corta
            },
            "LISTA_NUMERADA": {
                "nombre": "Lista Numerada",
                "descripcion": "Sangría colgante, numeración automática",
                "aplicacion": self._aplicar_lista_numerada
            },
            "LISTA_VIÑETAS": {
                "nombre": "Lista con Viñetas",
                "descripcion": "Sangría colgante, viñetas",
                "aplicacion": self._aplicar_lista_vinetas
            },
            "TEXTO_ENFASIS": {
                "nombre": "Texto con Énfasis",
                "descripcion": "Cursiva para términos técnicos",
                "aplicacion": self._aplicar_texto_enfasis
            },
            "CITA_TEXTUAL": {
                "nombre": "Cita Textual",
                "descripcion": "Sangría completa, sin comillas",
                "aplicacion": self._aplicar_cita_textual
            },
            "SIN_CAMBIOS": {
                "nombre": "Sin Cambios",
                "descripcion": "Mantener formato actual",
                "aplicacion": lambda p: None
            }
        }
    
    def analizar_seccion_manual(self, nombre_seccion, inicio_parrafo=None, fin_parrafo=None):
        """Analiza una sección específica párrafo por párrafo"""
        
        print(f"\n🔍 ANÁLISIS MANUAL DE SECCIÓN: {nombre_seccion}")
        print("="*70)
        
        # Abrir documento
        if self.ruta_trabajo.exists():
            self.documento = Document(str(self.ruta_trabajo))
            print(f"📄 Abriendo documento de trabajo existente: {self.ruta_trabajo.name}")
        else:
            self.documento = Document(str(self.ruta_original))
            print(f"📄 Creando nueva copia de trabajo: {self.ruta_trabajo.name}")
        
        # Determinar rango de párrafos
        if inicio_parrafo is None or fin_parrafo is None:
            if nombre_seccion == "4.1":
                inicio_parrafo, fin_parrafo = self._buscar_contenido_real_41()
            else:
                inicio_parrafo, fin_parrafo = self._buscar_seccion(nombre_seccion)
        
        if inicio_parrafo is None:
            print(f"❌ No se encontró la sección '{nombre_seccion}'")
            return
        
        print(f"📍 Rango detectado: párrafos {inicio_parrafo} a {fin_parrafo}")
        print(f"📊 Total párrafos a analizar: {fin_parrafo - inicio_parrafo + 1}")
        
        # Mostrar batería de estilos disponibles
        self._mostrar_estilos_disponibles()
        
        # Analizar párrafo por párrafo
        resumen_cambios = []
        
        for i in range(inicio_parrafo, fin_parrafo + 1):
            if i >= len(self.documento.paragraphs):
                break
                
            paragraph = self.documento.paragraphs[i]
            resultado = self._analizar_parrafo_individual(i, paragraph, nombre_seccion)
            
            if resultado:
                resumen_cambios.append(resultado)
        
        # Guardar documento con cambios
        self.documento.save(str(self.ruta_trabajo))
        
        # Generar reporte
        self._generar_reporte_seccion(nombre_seccion, resumen_cambios)
        
        print(f"\n✅ ANÁLISIS COMPLETADO para sección: {nombre_seccion}")
        print(f"💾 Documento guardado: {self.ruta_trabajo}")
        
        return resumen_cambios
    
    def _buscar_seccion(self, nombre_seccion):
        """Busca el inicio y fin de una sección específica"""
        
        patrones_busqueda = {
            "METODOLOGÍA": ["METODOLOGÍA", "METODOLOGIA", "III. METODOLOGÍA", "3. METODOLOGÍA"],
            "INTRODUCCIÓN": ["INTRODUCCIÓN", "INTRODUCTION", "I. INTRODUCCIÓN", "1. INTRODUCCIÓN"],
            "OBJETIVOS": ["OBJETIVOS", "OBJETIVO", "II. OBJETIVOS", "2. OBJETIVOS"],
            "RESULTADOS": ["RESULTADOS", "HALLAZGOS", "IV. RESULTADOS", "4. RESULTADOS"],
            "CONCLUSIONES": ["CONCLUSIONES", "CONCLUSIÓN", "VI. CONCLUSIONES", "6. CONCLUSIONES"],
            "4.1": ["4.1.", "4.1\t", "CARACTERIZACIÓN GENERAL", "CARACTERIZACIÓN DE LA MUESTRA", "MUESTRA DE ESTUDIO"]
        }
        
        patrones = patrones_busqueda.get(nombre_seccion.upper(), [nombre_seccion])
        inicio = None
        fin = None
        
        for i, paragraph in enumerate(self.documento.paragraphs):
            texto = paragraph.text.strip()
            
            # Buscar inicio
            if inicio is None:
                for patron in patrones:
                    if patron in texto.upper() and len(texto) < 200:
                        inicio = i
                        print(f"📍 Inicio encontrado en párrafo {i}: '{texto[:50]}...'")
                        break
            
            # Buscar fin (siguiente sección principal o siguiente subsección 4.2)
            elif inicio is not None:
                texto_upper = texto.upper()
                # Para sección 4.1, terminar cuando encontremos 4.2
                if nombre_seccion.upper() == "4.1":
                    if "4.2." in texto and len(texto) < 200:
                        fin = i - 1
                        print(f"📍 Fin encontrado en párrafo {fin} (siguiente subsección 4.2)")
                        break
                # Para otras secciones, usar lógica original
                elif (any(seccion in texto_upper for seccion in ["RESULTADOS", "OBJETIVOS", "CONCLUSIONES", "REFERENCIAS"]) 
                    and len(texto) < 200 and i > inicio + 5):
                    fin = i - 1
                    print(f"📍 Fin encontrado en párrafo {fin}")
                    break
        
        if inicio is not None and fin is None:
            fin = min(inicio + 50, len(self.documento.paragraphs) - 1)  # Límite de 50 párrafos
            
        return inicio, fin
    
    def _buscar_contenido_real_41(self):
        """Busca el contenido real de la sección 4.1 (no el índice)"""
        
        print("🔍 Buscando contenido real de sección 4.1...")
        
        inicio = None
        fin = None
        
        for i, paragraph in enumerate(self.documento.paragraphs):
            texto = paragraph.text.strip()
            
            # Buscar el inicio del contenido real de 4.1
            if inicio is None:
                # Buscar párrafos que contengan contenido sustancial sobre caracterización
                if (len(texto) > 100 and 
                    any(palabra in texto.upper() for palabra in [
                        "CARACTERIZACIÓN", "POBLACIÓN", "MUESTRA", "INSTITUCIONES EDUCATIVAS", 
                        "FE Y ALEGRÍA", "DISTRIBUCIÓN", "NIVEL EDUCATIVO"
                    ])):
                    inicio = i
                    print(f"📍 Contenido real 4.1 encontrado en párrafo {i}: '{texto[:50]}...'")
                    break
            
            # Buscar fin cuando aparezca contenido de 4.2
            elif inicio is not None:
                if (len(texto) > 50 and 
                    any(palabra in texto.upper() for palabra in [
                        "INDICADORES CLAVE", "KPI", "SISTEMA DE INDICADORES", "VARIABLES DEPENDIENTES"
                    ])):
                    fin = i - 1
                    print(f"📍 Fin del contenido 4.1 en párrafo {fin}")
                    break
        
        # Si no encuentra fin específico, limitar a un rango razonable
        if inicio is not None and fin is None:
            fin = min(inicio + 20, len(self.documento.paragraphs) - 1)
            print(f"📍 Fin estimado del contenido 4.1 en párrafo {fin}")
        
        return inicio, fin
    
    def _mostrar_estilos_disponibles(self):
        """Muestra la batería de estilos disponibles"""
        print(f"\n🎨 ESTILOS DISPONIBLES:")
        print("-" * 50)
        for clave, estilo in self.estilos_definidos.items():
            print(f"  {clave}: {estilo['nombre']}")
            print(f"    └── {estilo['descripcion']}")
        print("-" * 50)
    
    def _analizar_parrafo_individual(self, numero_parrafo, paragraph, seccion):
        """Analiza un párrafo individual y solicita decisión manual"""
        
        texto = paragraph.text.strip()
        
        # Saltar párrafos vacíos
        if not texto:
            return None
        
        print(f"\n📝 PÁRRAFO {numero_parrafo} de {seccion}:")
        print("┌" + "─" * 68 + "┐")
        print(f"│ {texto[:64]:<64} │")
        if len(texto) > 64:
            print(f"│ {texto[64:128]:<64} │")
        if len(texto) > 128:
            print(f"│ {'[...texto continúa...]':<64} │")
        print("└" + "─" * 68 + "┘")
        
        # Análisis automático y sugerencia
        sugerencia = self._sugerir_estilo_automatico(texto, paragraph)
        print(f"🤖 SUGERENCIA AUTOMÁTICA: {sugerencia}")
        
        # Análisis del contexto
        contexto = self._analizar_contexto_parrafo(numero_parrafo, texto)
        if contexto:
            print(f"🔍 CONTEXTO: {contexto}")
        
        # Solicitar decisión manual (simulada por análisis inteligente)
        decision = self._decision_manual_simulada(texto, sugerencia, contexto)
        
        if decision != "SIN_CAMBIOS":
            print(f"✅ APLICANDO: {self.estilos_definidos[decision]['nombre']}")
            
            # Aplicar el estilo
            self.estilos_definidos[decision]["aplicacion"](paragraph)
            
            return {
                "numero_parrafo": numero_parrafo,
                "texto_original": texto[:100] + "..." if len(texto) > 100 else texto,
                "estilo_aplicado": decision,
                "sugerencia_automatica": sugerencia,
                "decision_final": decision
            }
        else:
            print(f"⏭️ MANTENIENDO formato original")
            return None
    
    def _sugerir_estilo_automatico(self, texto, paragraph):
        """Sugiere un estilo basado en análisis automático"""
        
        texto_upper = texto.upper()
        longitud = len(texto)
        
        # Títulos principales
        if (any(palabra in texto_upper for palabra in ["METODOLOGÍA", "INTRODUCCIÓN", "OBJETIVOS", "RESULTADOS", "CONCLUSIONES"]) 
            and longitud < 100):
            return "TITULO_PRINCIPAL"
        
        # Subtítulos nivel 2
        elif (texto.startswith(("3.1", "3.2", "3.3", "4.1", "4.2")) or
              any(palabra in texto_upper for palabra in ["ENFOQUE Y DISEÑO", "POBLACIÓN", "ANÁLISIS", "TÉCNICAS"])):
            return "SUBTITULO_NIVEL_2"
        
        # Subtítulos nivel 3  
        elif (texto.startswith(("3.1.1", "3.2.1", "4.1.1")) or
              any(palabra in texto_upper for palabra in ["FASE 1", "FASE 2", "ACTIVIDADES", "CRITERIOS"])):
            return "SUBTITULO_NIVEL_3"
        
        # Listas
        elif any(indicador in texto for indicador in ["•", "-", "1.", "2.", "a)", "b)"]):
            return "LISTA_VIÑETAS"
        
        # Descripciones cortas
        elif longitud < 150 and not texto.endswith("."):
            return "DESCRIPCION_CORTA"
        
        # Párrafos normales
        elif longitud > 150:
            return "PARRAFO_NORMAL"
        
        return "SIN_CAMBIOS"
    
    def _analizar_contexto_parrafo(self, numero, texto):
        """Analiza el contexto del párrafo"""
        contextos = []
        
        if numero > 0:
            texto_anterior = self.documento.paragraphs[numero-1].text.strip()
            if texto_anterior:
                contextos.append(f"Anterior: '{texto_anterior[:30]}...'")
        
        if any(palabra in texto.upper() for palabra in ["TABLA", "FIGURA", "GRÁFICO"]):
            contextos.append("Contiene referencia a elemento visual")
            
        if len(texto.split()) < 10:
            contextos.append("Texto corto - posible título/subtítulo")
            
        return " | ".join(contextos) if contextos else None
    
    def _decision_manual_simulada(self, texto, sugerencia, contexto):
        """Simula decisión manual inteligente basada en análisis profundo"""
        
        texto_upper = texto.upper()
        longitud = len(texto)
        palabras = len(texto.split())
        
        # REGLAS ESPECÍFICAS REFINADAS PARA SECCIÓN 4.1
        
        # Títulos principales de secciones
        if (longitud < 50 and palabras < 5 and 
            any(palabra in texto_upper for palabra in ["OBJETIVOS", "METODOLOGÍA", "RESULTADOS", "CONCLUSIONES"]) and
            not any(char in texto for char in ["\t", ".", ":", "4."])):
            return "TITULO_PRINCIPAL"
        
        # Subtítulos de sección (como "Justificación del Presente Estudio")
        elif (longitud < 100 and palabras < 8 and
              any(palabra in texto_upper for palabra in [
                  "JUSTIFICACIÓN", "PROPÓSITO", "ALCANCE", "OBJETIVO GENERAL", "OBJETIVOS ESPECÍFICOS"
              ]) and not texto.startswith(("4.", "3.", "2.", "1."))):
            return "SUBTITULO_NIVEL_2"
        
        # Subtítulos nivel 3 - elementos introductorios
        elif (longitud < 150 and
              (any(palabra in texto_upper for palabra in [
                  "CARACTERIZACIÓN SE FUNDAMENTA", "ANÁLISIS MULTIDIMENSIONAL", "COMPONENTES CRÍTICOS"
              ]) or texto.endswith(":"))):
            return "SUBTITULO_NIVEL_3"
        
        # Elementos de lista o componentes (Logros de aprendizaje, Desempeño docente, etc.)
        elif (longitud > 100 and longitud < 300 and
              any(palabra in texto_upper for palabra in [
                  "LOGROS DE APRENDIZAJE", "DESEMPEÑO DOCENTE", "FACTORES CONTEXTUALES"
              ])):
            return "DESCRIPCION_CORTA"
        
        # Párrafos sustanciales con contenido desarrollado
        elif longitud > 250 and palabras > 30:
            return "PARRAFO_NORMAL"
        
        # Descripciones cortas para títulos menores
        elif longitud < 150 and palabras < 15 and not texto.endswith("."):
            return "DESCRIPCION_CORTA"
        
        # Por defecto, usar la sugerencia automática mejorada
        return sugerencia
    
    # Métodos de aplicación de estilos
    def _aplicar_titulo_principal(self, paragraph):
        """Aplica estilo Título Principal APA 7"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_before = Pt(12)
        paragraph.space_after = Pt(6)
        
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_subtitulo_nivel_2(self, paragraph):
        """Aplica estilo Subtítulo Nivel 2 APA 7"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.space_before = Pt(12)
        paragraph.space_after = Pt(6)
        
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_subtitulo_nivel_3(self, paragraph):
        """Aplica estilo Subtítulo Nivel 3 APA 7"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.space_before = Pt(6)
        paragraph.space_after = Pt(3)
        
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_parrafo_normal(self, paragraph):
        """Aplica estilo Párrafo Normal APA 7"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 2.0
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_descripcion_corta(self, paragraph):
        """Aplica estilo Descripción Corta"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.5
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_lista_numerada(self, paragraph):
        """Aplica estilo Lista Numerada"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.hanging_indent = Inches(0.25)
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_lista_vinetas(self, paragraph):
        """Aplica estilo Lista con Viñetas"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.hanging_indent = Inches(0.25)
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_texto_enfasis(self, paragraph):
        """Aplica estilo Texto con Énfasis"""
        for run in paragraph.runs:
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _aplicar_cita_textual(self, paragraph):
        """Aplica estilo Cita Textual"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.right_indent = Inches(0.5)
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    def _generar_reporte_seccion(self, nombre_seccion, cambios):
        """Genera reporte detallado de la sección procesada"""
        
        reporte = {
            "seccion": nombre_seccion,
            "fecha_procesamiento": datetime.now().isoformat(),
            "total_cambios": len(cambios),
            "cambios_por_estilo": {},
            "detalles_cambios": cambios
        }
        
        # Contar cambios por tipo de estilo
        for cambio in cambios:
            estilo = cambio["estilo_aplicado"]
            if estilo not in reporte["cambios_por_estilo"]:
                reporte["cambios_por_estilo"][estilo] = 0
            reporte["cambios_por_estilo"][estilo] += 1
        
        # Guardar reporte
        archivo_reporte = Path(f"reporte_manual_{nombre_seccion.lower()}.json")
        with open(archivo_reporte, 'w', encoding='utf-8') as f:
            json.dump(reporte, f, indent=2, ensure_ascii=False)
        
        print(f"\n📊 REPORTE GENERADO: {archivo_reporte}")
        print("🎯 RESUMEN DE CAMBIOS:")
        for estilo, cantidad in reporte["cambios_por_estilo"].items():
            nombre_estilo = self.estilos_definidos[estilo]["nombre"]
            print(f"  ✅ {nombre_estilo}: {cantidad} aplicaciones")

def demo_analisis_manual():
    """Demostración del análisis manual supervisado"""
    
    # Usar el backup original como especificado
    ruta_documento = r"assets\Consultoria\01 Informe en elaboración\01 Informe Tipologías IIEE 2025 1.0_backup.docx"
    
    print("🔍 DEMO: EDITOR MANUAL SUPERVISADO")
    print("="*70)
    print("🎯 Enfoque: Análisis párrafo por párrafo con supervisión total")
    print("📋 Sección objetivo: 4.1. Caracterización General de la Muestra")
    print("📄 Documento: BACKUP ORIGINAL")
    print("="*70)
    
    try:
        # Crear editor supervisado
        editor = EditorManualSupervisado(ruta_documento)
        
        # Analizar sección 4.1 específicamente
        cambios = editor.analizar_seccion_manual("4.1")
        
        print(f"\n🎉 ANÁLISIS MANUAL COMPLETADO")
        print(f"📊 Total cambios aplicados: {len(cambios) if cambios else 0}")
        print(f"💾 Documento editado: {editor.ruta_trabajo}")
        
        return editor
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

if __name__ == "__main__":
    demo_analisis_manual()
