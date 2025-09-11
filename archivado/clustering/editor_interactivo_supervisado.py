#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EDITOR INTERACTIVO SUPERVISADO APA 7
Versión completamente manual con input del usuario párrafo por párrafo

Permite al usuario tomar decisiones manuales sobre cada párrafo
"""

from editor_manual_supervisado import EditorManualSupervisado
import json
from pathlib import Path

class EditorInteractivoSupervisado(EditorManualSupervisado):
    """Editor con supervisión manual real del usuario"""
    
    def analizar_parrafo_interactivo(self, numero_parrafo, paragraph, seccion):
        """Análisis interactivo con input del usuario"""
        
        texto = paragraph.text.strip()
        
        # Saltar párrafos vacíos
        if not texto:
            return None
        
        print(f"\n📝 PÁRRAFO {numero_parrafo} de {seccion}:")
        print("┌" + "─" * 78 + "┐")
        
        # Mostrar texto completo del párrafo
        lines = [texto[i:i+76] for i in range(0, len(texto), 76)]
        for line in lines[:3]:  # Máximo 3 líneas
            print(f"│ {line:<76} │")
        if len(lines) > 3:
            print(f"│ {'[...continúa, ' + str(len(texto)) + ' caracteres total]':<76} │")
            
        print("└" + "─" * 78 + "┘")
        
        # Análisis automático
        sugerencia = self._sugerir_estilo_automatico(texto, paragraph)
        contexto = self._analizar_contexto_parrafo(numero_parrafo, texto)
        
        print(f"🤖 SUGERENCIA AUTOMÁTICA: {sugerencia}")
        if contexto:
            print(f"🔍 CONTEXTO: {contexto}")
        
        # Mostrar opciones disponibles
        print(f"\n🎨 OPCIONES DE ESTILO:")
        opciones = list(self.estilos_definidos.keys())
        for i, clave in enumerate(opciones, 1):
            estilo = self.estilos_definidos[clave]
            print(f"  {i}. {clave}: {estilo['nombre']}")
        
        # Solicitar input del usuario
        while True:
            try:
                print(f"\n❓ ¿Qué estilo aplicar? (1-{len(opciones)}, 's' para saltar, 'q' para salir):")
                respuesta = input(">>> ").strip().lower()
                
                if respuesta == 'q':
                    print("🛑 Saliendo del editor...")
                    return "QUIT"
                elif respuesta == 's':
                    print("⏭️ Saltando párrafo sin cambios")
                    return None
                else:
                    indice = int(respuesta) - 1
                    if 0 <= indice < len(opciones):
                        decision = opciones[indice]
                        break
                    else:
                        print(f"❌ Opción inválida. Ingrese 1-{len(opciones)}")
                        
            except ValueError:
                print("❌ Ingrese un número válido")
        
        # Confirmar decisión
        estilo_elegido = self.estilos_definidos[decision]
        print(f"\n✅ ELEGIDO: {estilo_elegido['nombre']}")
        print(f"   └── {estilo_elegido['descripcion']}")
        
        confirmacion = input("¿Confirmar aplicación? (y/n): ").strip().lower()
        
        if confirmacion == 'y':
            # Aplicar el estilo
            estilo_elegido["aplicacion"](paragraph)
            print(f"✅ ESTILO APLICADO: {estilo_elegido['nombre']}")
            
            return {
                "numero_parrafo": numero_parrafo,
                "texto_original": texto[:100] + "..." if len(texto) > 100 else texto,
                "estilo_aplicado": decision,
                "sugerencia_automatica": sugerencia,
                "decision_manual": decision
            }
        else:
            print("❌ Aplicación cancelada")
            return None
    
    def analizar_seccion_interactiva(self, nombre_seccion, inicio_parrafo=None, fin_parrafo=None):
        """Análisis interactivo completo de una sección"""
        
        print(f"\n🔍 ANÁLISIS INTERACTIVO DE SECCIÓN: {nombre_seccion}")
        print("="*80)
        
        # Preparar documento
        if self.ruta_trabajo.exists():
            self.documento = Document(str(self.ruta_trabajo))
            print(f"📄 Continuando con: {self.ruta_trabajo.name}")
        else:
            self.documento = Document(str(self.ruta_original))
            print(f"📄 Creando nueva copia: {self.ruta_trabajo.name}")
        
        # Determinar rango
        if inicio_parrafo is None or fin_parrafo is None:
            inicio_parrafo, fin_parrafo = self._buscar_seccion(nombre_seccion)
        
        if inicio_parrafo is None:
            print(f"❌ No se encontró la sección '{nombre_seccion}'")
            return
        
        print(f"📍 Rango: párrafos {inicio_parrafo} a {fin_parrafo}")
        print(f"📊 Total: {fin_parrafo - inicio_parrafo + 1} párrafos")
        
        # Mostrar instrucciones
        print(f"\n📋 INSTRUCCIONES:")
        print(f"  • Revise cada párrafo cuidadosamente")
        print(f"  • Elija el estilo más apropiado según contenido")
        print(f"  • Use 's' para saltar sin cambios")
        print(f"  • Use 'q' para salir y guardar progreso")
        print(f"  • Presione Enter después de cada decisión")
        
        input(f"\n🚀 Presione Enter para comenzar...")
        
        # Procesar párrafos interactivamente
        cambios = []
        
        for i in range(inicio_parrafo, fin_parrafo + 1):
            if i >= len(self.documento.paragraphs):
                break
                
            paragraph = self.documento.paragraphs[i]
            resultado = self.analizar_parrafo_interactivo(i, paragraph, nombre_seccion)
            
            if resultado == "QUIT":
                print(f"\n🛑 Proceso interrumpido por el usuario")
                break
            elif resultado:
                cambios.append(resultado)
        
        # Guardar progreso
        self.documento.save(str(self.ruta_trabajo))
        
        # Generar reporte
        self._generar_reporte_interactivo(nombre_seccion, cambios)
        
        print(f"\n✅ SESIÓN COMPLETADA")
        print(f"📊 Cambios aplicados: {len(cambios)}")
        print(f"💾 Documento guardado: {self.ruta_trabajo}")
        
        return cambios
    
    def _generar_reporte_interactivo(self, seccion, cambios):
        """Genera reporte de sesión interactiva"""
        
        reporte = {
            "tipo_sesion": "interactiva_supervisada",
            "seccion": seccion,
            "fecha": datetime.now().isoformat(),
            "total_cambios": len(cambios),
            "estadisticas": {},
            "detalles": cambios
        }
        
        # Estadísticas por estilo
        for cambio in cambios:
            estilo = cambio["estilo_aplicado"]
            if estilo not in reporte["estadisticas"]:
                reporte["estadisticas"][estilo] = 0
            reporte["estadisticas"][estilo] += 1
        
        # Guardar reporte
        archivo = Path(f"sesion_interactiva_{seccion.lower()}.json")
        with open(archivo, 'w', encoding='utf-8') as f:
            json.dump(reporte, f, indent=2, ensure_ascii=False)
        
        print(f"📊 Reporte guardado: {archivo}")

def demo_sesion_interactiva():
    """Demo de sesión completamente interactiva"""
    
    ruta_documento = r"assets\Consultoria\01 Informe en elaboración\01 Informe Tipologías IIEE 2025 1.0.docx"
    
    print("🎯 EDITOR INTERACTIVO SUPERVISADO")
    print("="*50)
    print("🔍 Control manual total párrafo por párrafo")
    print("📋 Sección: METODOLOGÍA")
    print("="*50)
    
    try:
        from docx import Document
        from datetime import datetime
        
        # Crear editor interactivo
        editor = EditorInteractivoSupervisado(ruta_documento)
        
        # Confirmar inicio
        continuar = input("¿Iniciar análisis interactivo? (y/n): ").strip().lower()
        
        if continuar == 'y':
            # Analizar metodología de forma interactiva
            cambios = editor.analizar_seccion_interactiva("METODOLOGÍA")
            
            print(f"\n🎉 SESIÓN INTERACTIVA FINALIZADA")
            if cambios:
                print(f"📊 Total de cambios: {len(cambios)}")
                print(f"💾 Documento editado: {editor.ruta_trabajo}")
            else:
                print("📝 No se realizaron cambios")
        else:
            print("❌ Sesión cancelada")
            
        return editor
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        return None

if __name__ == "__main__":
    demo_sesion_interactiva()
