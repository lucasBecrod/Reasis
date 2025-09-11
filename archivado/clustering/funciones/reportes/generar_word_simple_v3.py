#!/usr/bin/env python3
"""
Generador de Documento Word Ejecutivo Simplificado - Tipologías v3.1
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_executive_document():
    """Crea el documento Word ejecutivo"""
    
    doc = Document()
    
    # Título principal
    title = doc.add_heading('TIPOLOGÍAS INSTITUCIONALES FE Y ALEGRÍA 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Análisis Definitivo con 6 Tipologías Diferenciadas')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    
    # Información del proyecto
    doc.add_paragraph('Fecha: 12 de Agosto de 2025')
    doc.add_paragraph('Proyecto: Reasis - Tipologías de Instituciones Educativas Fe y Alegría')
    doc.add_paragraph('Metodología: Clustering K-Means Avanzado con 21+ Variables')
    
    doc.add_paragraph()
    
    # RESUMEN EJECUTIVO
    doc.add_heading("RESUMEN EJECUTIVO", 1)
    
    doc.add_heading("Objetivo Conseguido:", level=2)
    p = doc.add_paragraph()
    p.add_run("✅ Generación exitosa de 6 tipologías institucionales diferenciadas ").bold = True
    p.add_run("mediante análisis científico robusto con 184 instituciones, 21+ variables y validación multi-métrica.")
    
    doc.add_heading("Alcance del Análisis:", level=2)
    alcance_items = [
        "184 instituciones educativas clasificadas (48% del total Fe y Alegría)",
        "21+ variables metodológicas y contextuales estandarizadas", 
        "6 tipologías específicas identificadas con alta coherencia interna",
        "6 redes territoriales completamente representadas",
        "Validación científica mediante múltiples algoritmos y métricas"
    ]
    
    for item in alcance_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # TIPOLOGÍAS IDENTIFICADAS
    doc.add_heading("TIPOLOGÍAS IDENTIFICADAS", 1)
    
    # Tipología 0
    doc.add_heading("TIPOLOGÍA 0: INSTITUCIONES ESPECIALIZADAS INICIALES", level=2)
    doc.add_paragraph("4 instituciones (2.2% del análisis)").runs[0].bold = True
    
    doc.add_heading("Perfil Característico:", level=3)
    perfil_0 = [
        "🎓 Nivel Educativo: Inicial/No Escolarizada (diferenciador clave)",
        "💻 Competencia Digital: Muy Alta (liderazgo tecnológico)", 
        "🌐 Gestión: Especializada con enfoques diferenciados",
        "👥 Ratio Estudiante-Docente: Muy Bajo (atención personalizada)",
        "📍 Concentración: Red 47 principalmente"
    ]
    for item in perfil_0:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Estrategias Recomendadas:", level=3)
    estrategias_0 = [
        "Consolidar especialización en modelo pedagógico inicial diferenciado",
        "Establecer liderazgo tecnológico como referente digital",
        "Implementar mentoría compartiendo buenas prácticas de atención personalizada",
        "Evaluar replicabilidad del modelo en otras instituciones"
    ]
    for estrategia in estrategias_0:
        doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_page_break()
    
    # Tipología 1  
    doc.add_heading("TIPOLOGÍA 1: INSTITUCIONES RESILIENTES RURALES", level=2)
    doc.add_paragraph("36 instituciones (19.6% del análisis)").runs[0].bold = True
    
    doc.add_heading("Perfil Característico:", level=3)
    perfil_1 = [
        "🌾 Contexto: Altamente Rural con excelente adaptación",
        "📚 Logro Académico: Alto Relativo (resiliencia excepcional)",
        "⚡ Fortaleza: Excelente rendimiento pese al contexto desafiante",
        "👨‍🏫 Desempeño Docente: Promedio Alto con compromiso territorial",
        "🔧 Organización: Multigrado/Unidocente optimizada y efectiva"
    ]
    for item in perfil_1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Distribución Territorial:", level=3)
    distribucion_1 = [
        "Red 47: 18 instituciones (50.0%) - Concentración principal",
        "Red 79: 12 instituciones (33.3%) - Segundo grupo importante",
        "Red 48: 4 instituciones (11.1%)",
        "Red 72: 2 instituciones (5.6%)"
    ]
    for item in distribucion_1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Estrategias Recomendadas:", level=3)
    estrategias_1 = [
        "Reconocer y visibilizar logros excepcionales en contexto rural",
        "Fortalecer y sostener factores de éxito identificados",
        "Crear redes de intercambio entre instituciones rurales exitosas", 
        "Sistematizar y documentar metodologías rurales efectivas"
    ]
    for estrategia in estrategias_1:
        doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_page_break()
    
    # Tipología 2
    doc.add_heading("TIPOLOGÍA 2: INSTITUCIONES DE ALTO RENDIMIENTO", level=2)
    doc.add_paragraph("63 instituciones (34.2% del análisis) - GRUPO PRINCIPAL").runs[0].bold = True
    
    doc.add_heading("Perfil Característico:", level=3)
    perfil_2 = [
        "📈 Rendimiento: Consistentemente Alto en múltiples dimensiones",
        "💪 Fortalezas: Equilibrio excepcional en todas las variables clave",
        "👥 Ratio Óptimo: Estudiante-Docente perfectamente balanceado",
        "🎯 Gestión: Sólida y sostenible con resultados probados",
        "🌟 Liderazgo: Referentes indiscutibles en la red Fe y Alegría"
    ]
    for item in perfil_2:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Distribución Territorial:", level=3)
    distribucion_2 = [
        "Red 44: 22 instituciones (34.9%) - Liderazgo urbano",
        "Red 48: 16 instituciones (25.4%) - Equilibrio mixto", 
        "Red 72: 12 instituciones (19.0%) - Consolidación regional",
        "Red 54: 8 instituciones (12.7%)",
        "Red 79: 5 instituciones (7.9%)"
    ]
    for item in distribucion_2:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Estrategias Recomendadas:", level=3)
    estrategias_2 = [
        "Asumir rol de liderazgo y mentoría hacia otras tipologías",
        "Implementar procesos de mejora continua sistemática",
        "Sistematizar y documentar buenas prácticas replicables",
        "Liderar pilotaje de nuevas metodologías innovadoras"
    ]
    for estrategia in estrategias_2:
        doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_page_break()
    
    # Tipología 3
    doc.add_heading("TIPOLOGÍA 3: INSTITUCIONES EQUILIBRADAS", level=2)
    doc.add_paragraph("19 instituciones (10.3% del análisis)").runs[0].bold = True
    
    doc.add_heading("Perfil Característico:", level=3)
    perfil_3 = [
        "⚖️ Balance: Promedio balanceado en todas las dimensiones",
        "📊 Estabilidad: Resultados consistentes sin extremos",
        "🎯 Potencial: Oportunidades claras y bien definidas de mejoramiento",
        "👨‍🏫 Docentes: Desempeño promedio con margen de crecimiento",
        "📈 Tendencia: Progreso sostenido y gradual comprobado"
    ]
    for item in perfil_3:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Estrategias Recomendadas:", level=3)
    estrategias_3 = [
        "Implementar intervención focalizada en 2-3 áreas prioritarias por institución",
        "Desarrollar estrategias de mejora gradual y sostenida",
        "Establecer acompañamiento y mentoría desde Tipología 2",
        "Implementar monitoreo estrecho de indicadores clave de progreso"
    ]
    for estrategia in estrategias_3:
        doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_page_break()
    
    # Tipología 4
    doc.add_heading("TIPOLOGÍA 4: INSTITUCIONES EN DESARROLLO", level=2)
    doc.add_paragraph("51 instituciones (27.7% del análisis) - SEGUNDO GRUPO PRINCIPAL").runs[0].bold = True
    
    doc.add_heading("Perfil Característico:", level=3)
    perfil_4 = [
        "🔧 Necesidades: Desarrollo prioritario en competencias docentes y digitales",
        "📱 Tecnología: Brecha digital significativa que requiere atención",
        "👨‍🏫 Capacitación: Docentes requieren fortalecimiento específico",
        "🎯 Oportunidad: Alto potencial de mejora con intervención adecuada",
        "📍 Contexto: Mixto rural-urbano con desafíos diferenciados"
    ]
    for item in perfil_4:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Estrategias Recomendadas:", level=3)
    estrategias_4 = [
        "Implementar fortalecimiento docente intensivo en competencias clave identificadas",
        "Desarrollar programa específico y sistemático contra la brecha digital",
        "Mejorar infraestructura tecnológica y conectividad institucional",
        "Focalizar recursos en 3-4 competencias prioritarias por institución"
    ]
    for estrategia in estrategias_4:
        doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_page_break()
    
    # Tipología 5
    doc.add_heading("TIPOLOGÍA 5: INSTITUCIONES URBANAS COMPLEJAS", level=2)
    doc.add_paragraph("11 instituciones (6.0% del análisis)").runs[0].bold = True
    
    doc.add_heading("Perfil Característico:", level=3)
    perfil_5 = [
        "🏙️ Contexto: Urbano de alta densidad poblacional", 
        "👥 Complejidad: Ratio estudiante-docente elevado (gestión de escala)",
        "🎯 Desafíos: Gestión de escala y atención diferenciada compleja",
        "📊 Rendimiento: Variable según dimensión específica analizada",
        "🔧 Especialización: Requerimientos urbanos específicos y únicos"
    ]
    for item in perfil_5:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Estrategias Recomendadas:", level=3)
    estrategias_5 = [
        "Optimizar gestión organizacional para contextos de alta densidad",
        "Desarrollar estrategias específicas de atención diferenciada para ratios elevados", 
        "Implementar metodologías especializadas para contexto urbano complejo",
        "Coordinar con Tipología 2 para transferir modelos urbanos exitosos"
    ]
    for estrategia in estrategias_5:
        doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_page_break()
    
    # VARIABLES DISCRIMINANTES
    doc.add_heading("FACTORES DIFERENCIADORES CLAVE", 1)
    
    doc.add_paragraph("Las siguientes variables muestran las mayores diferencias entre tipologías:")
    
    discriminantes = [
        "🌾 Ruralidad vs Urbanidad: Factor diferenciador principal entre contextos geográficos",
        "👥 Ratio Estudiante-Docente: Discrimina entre atención personalizada y gestión de escala", 
        "👨‍🏫 Desempeño Docente: Separa instituciones con fortaleza vs. necesidad de capacitación",
        "💻 Competencia Digital: Diferencia instituciones tecnológicamente avanzadas de las rezagadas"
    ]
    
    for item in discriminantes:
        doc.add_paragraph(item, style='List Bullet')
    
    # CONCLUSIONES
    doc.add_heading("CONCLUSIONES Y PRÓXIMOS PASOS", 1)
    
    doc.add_heading("Logros Conseguidos:", level=2)
    logros = [
        "✅ 6 tipologías institucionales robustas identificadas con base científica sólida",
        "✅ Metodología avanzada con 21+ variables y validación cruzada múltiple",
        "✅ Estrategias específicas y directamente aplicables para cada tipología",
        "✅ Red de 63 instituciones líderes preparadas para mentoría y transferencia",
        "✅ Base empírica robusta para toma de decisiones diferenciadas por contexto"
    ]
    
    for logro in logros:
        doc.add_paragraph(logro, style='List Bullet')
    
    doc.add_heading("Recomendaciones de Implementación Inmediata:", level=2)
    recomendaciones = [
        "🔄 Validación participativa con equipos territoriales Fe y Alegría (30 días)",
        "🎯 Diseño detallado de intervenciones específicas por cada tipología (30 días)",
        "🚀 Pilotaje controlado en instituciones representativas de cada tipo (30 días)",
        "📊 Implementación de sistema integral de monitoreo continuo (90 días)",
        "🤝 Establecimiento de red de mentoría inter-tipologías (60 días)"
    ]
    
    for rec in recomendaciones:
        doc.add_paragraph(rec, style='List Bullet')
    
    doc.add_heading("Impacto Esperado:", level=2)
    impactos = [
        "🎯 Intervenciones educativas más precisas y efectivas por contexto específico",
        "📈 Optimización de recursos mediante focalización estratégica diferenciada",
        "🌟 Fortalecimiento de red de liderazgo institucional Fe y Alegría",
        "📊 Sistema de gestión educativa basado en evidencia científica robusta",
        "🚀 Modelo replicable para otras redes educativas a nivel nacional"
    ]
    
    for impacto in impactos:
        doc.add_paragraph(impacto, style='List Bullet')
    
    # Pie de página
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph("PROYECTO REASIS - TIPOLOGÍAS INSTITUCIONALES FE Y ALEGRÍA 2025")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].bold = True
    
    subfooter = doc.add_paragraph("184 instituciones analizadas | 6 tipologías diferenciadas | 21+ variables | Metodología científica avanzada")
    subfooter.alignment = WD_ALIGN_PARAGRAPH.CENTER  
    subfooter.runs[0].font.size = Pt(9)
    subfooter.runs[0].italic = True
    
    return doc

def main():
    """Función principal"""
    
    try:
        # Crear documento
        doc = create_executive_document()
        
        # Guardar en ubicación apropiada
        output_path = Path("assets/Consultoria/01 Informe en elaboración/INFORME_EJECUTIVO_TIPOLOGIAS_V3_FINAL.docx")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(str(output_path))
        print(f"[OK] Documento Word ejecutivo generado: {output_path}")
        
        # También guardar una copia en la carpeta de reportes
        backup_path = Path("data/reports/INFORME_EJECUTIVO_TIPOLOGIAS_V3_FINAL.docx")
        backup_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(backup_path))
        print(f"[OK] Copia de respaldo generada: {backup_path}")
        
        return True
        
    except Exception as e:
        print(f"[ERROR] Error al generar documento: {e}")
        return False

if __name__ == "__main__":
    main()