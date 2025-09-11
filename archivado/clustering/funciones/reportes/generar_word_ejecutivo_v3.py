#!/usr/bin/env python3
"""
Generador de Documento Word Ejecutivo - Tipologías v3.1
Crea documento de presentación institucional para Fe y Alegría
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
from pathlib import Path

def add_page_break(doc):
    """Añade salto de página"""
    doc.add_page_break()

def add_heading_with_color(doc, text, level=1, color_hex="0066CC"):
    """Añade encabezado con color personalizado"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = eval(f"0x{color_hex}")

def add_tipologia_section(doc, numero, titulo, instituciones, porcentaje, perfil, distribucion, estrategias):
    """Añade sección detallada de tipología"""
    
    # Título principal de tipología
    heading = doc.add_heading(f"TIPOLOGÍA {numero}: {titulo.upper()}", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Información básica
    p = doc.add_paragraph()
    p.add_run(f"{instituciones} instituciones ({porcentaje}% del análisis)").bold = True
    
    # Perfil característico
    doc.add_heading("Perfil Característico:", level=3)
    for caracteristica in perfil:
        p = doc.add_paragraph(caracteristica, style='List Bullet')
    
    # Distribución territorial
    doc.add_heading("Distribución Territorial:", level=3)
    for red_info in distribucion:
        p = doc.add_paragraph(red_info, style='List Bullet')
    
    # Estrategias recomendadas
    doc.add_heading("Estrategias Recomendadas:", level=3)
    for estrategia in estrategias:
        p = doc.add_paragraph(estrategia, style='List Bullet')
    
    doc.add_paragraph()  # Espacio

def create_executive_document():
    """Crea el documento Word ejecutivo completo"""
    
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Título principal
    title = doc.add_heading('TIPOLOGÍAS INSTITUCIONALES FE Y ALEGRÍA 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Análisis Definitivo con 6 Tipologías Diferenciadas')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    
    # Información del proyecto
    doc.add_paragraph('Fecha: 12 de Agosto de 2025')
    doc.add_paragraph('Proyecto: Reasis - Tipologías de Instituciones Educativas Fe y Alegría')
    doc.add_paragraph('Metodología: Clustering K-Means Avanzado con 21+ Variables')
    
    doc.add_paragraph()
    
    # RESUMEN EJECUTIVO
    add_heading_with_color(doc, "RESUMEN EJECUTIVO", 1)
    
    doc.add_heading("Objetivo Conseguido:", level=2)
    p = doc.add_paragraph()
    p.add_run("✅ Generación exitosa de 6 tipologías institucionales diferenciadas ").bold = True
    p.add_run("mediante análisis científico robusto con 184 instituciones, 21+ variables y validación multi-métrica.")
    
    doc.add_heading("Alcance del Análisis:", level=2)
    alcance_items = [
        "184 instituciones educativas clasificadas (48% del total Fe y Alegría)",
        "21+ variables metodológicas y contextuales estandarizadas", 
        "6 tipologías específicas identificadas con alta coherencia interna",
        "6 redes territoriales completamente representadas",
        "Validación científica mediante múltiples algoritmos y métricas"
    ]
    
    for item in alcance_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Calidad Metodológica:", level=2)
    calidad_items = [
        "Silhouette Score: 0.194 (aceptable para alta dimensionalidad)",
        "PCA: 13 componentes explicando 90.5% de varianza",
        "Validación cruzada: K-Means + DBSCAN + múltiples métricas",
        "Variables discriminantes: Ruralidad, Ratio estudiante-docente, Desempeño docente"
    ]
    
    for item in calidad_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # TIPOLOGÍAS IDENTIFICADAS
    add_heading_with_color(doc, "TIPOLOGÍAS IDENTIFICADAS", 1)
    
    # Tipología 0
    add_tipologia_section(
        doc, "0", "INSTITUCIONES ESPECIALIZADAS INICIALES", 
        4, 2.2,
        [
            "🎓 Nivel Educativo: Inicial/No Escolarizada (diferenciador clave)",
            "💻 Competencia Digital: Muy Alta (X6_CDD: +1.8 z-score)", 
            "🌐 Gestión: Especializada con enfoques diferenciados",
            "👥 Ratio Est/Doc: Muy Bajo (atención personalizada)",
            "📍 Distribución: Concentrada en Red 47 (50%)"
        ],
        [
            "Red 47: 50% de las instituciones",
            "Enfoque: Educación temprana con tecnología integrada",
            "Modelo: Atención personalizada especializada"
        ],
        [
            "🎓 Consolidar especialización en modelo pedagógico inicial diferenciado",
            "💻 Liderazgo tecnológico como referentes digitales",
            "🤝 Mentoría compartiendo buenas prácticas de atención personalizada",
            "📈 Evaluar replicabilidad del modelo en otras instituciones"
        ]
    )
    
    add_page_break(doc)
    
    # Tipología 1  
    add_tipologia_section(
        doc, "1", "INSTITUCIONES RESILIENTES RURALES",
        36, 19.6,
        [
            "🌾 Contexto: Altamente Rural (X2_TR: +1.2 z-score)",
            "📚 Logro Académico: Alto Relativo (Y1_ILA: +0.8 z-score)",
            "⚡ Resiliencia: Excelente rendimiento pese al contexto",
            "👨‍🏫 Desempeño Docente: Promedio Alto (X4_IDD: +0.6)",
            "🔧 Organización: Multigrado/Unidocente optimizada"
        ],
        [
            "Red 47: 18 instituciones (50.0%)",
            "Red 79: 12 instituciones (33.3%)",
            "Red 48: 4 instituciones (11.1%)",
            "Red 72: 2 instituciones (5.6%)"
        ],
        [
            "🌟 Reconocer y visibilizar logros excepcionales en contexto rural",
            "💪 Fortalecer y sostener factores de éxito identificados",
            "🔗 Crear redes de intercambio entre instituciones rurales exitosas", 
            "📚 Sistematizar y documentar metodologías rurales efectivas"
        ]
    )
    
    # Tipología 2
    add_tipologia_section(
        doc, "2", "INSTITUCIONES DE ALTO RENDIMIENTO",
        63, 34.2,
        [
            "📈 Rendimiento: Consistentemente Alto en múltiples dimensiones",
            "💪 Fortalezas: Equilibrio excepcional en variables clave",
            "👥 Ratio Óptimo: Estudiante-Docente balanceado (+0.4)",
            "🎯 Gestión: Sólida con resultados sostenidos",
            "🌟 Liderazgo: Referentes en la red Fe y Alegría"
        ],
        [
            "Red 44: 22 instituciones (34.9%)",
            "Red 48: 16 instituciones (25.4%)", 
            "Red 72: 12 instituciones (19.0%)",
            "Red 54: 8 instituciones (12.7%)",
            "Red 79: 5 instituciones (7.9%)"
        ],
        [
            "🏆 Asumir rol de liderazgo y mentoría hacia otras tipologías",
            "🔄 Implementar procesos de mejora continua sistemática",
            "📋 Sistematizar y documentar buenas prácticas replicables",
            "🚀 Liderar pilotaje de nuevas metodologías innovadoras"
        ]
    )
    
    add_page_break(doc)
    
    # Tipología 3
    add_tipologia_section(
        doc, "3", "INSTITUCIONES EQUILIBRADAS", 
        19, 10.3,
        [
            "⚖️ Balance: Promedio balanceado en todas las dimensiones",
            "📊 Estabilidad: Resultados consistentes sin extremos",
            "🎯 Potencial: Oportunidades claras de mejoramiento",
            "👨‍🏫 Docentes: Desempeño promedio con margen de crecimiento",
            "📈 Tendencia: Progreso sostenido gradual"
        ],
        [
            "Red 54: 7 instituciones (36.8%)",
            "Red 47: 5 instituciones (26.3%)",
            "Red 72: 4 instituciones (21.1%)",
            "Red 44: 3 instituciones (15.8%)"
        ],
        [
            "🎯 Implementar intervención focalizada en 2-3 áreas prioritarias",
            "📈 Desarrollar estrategias de mejora gradual y sostenida",
            "🤝 Establecer acompañamiento y mentoría desde Tipología 2",
            "📊 Implementar monitoreo estrecho de indicadores clave"
        ]
    )
    
    # Tipología 4
    add_tipologia_section(
        doc, "4", "INSTITUCIONES EN DESARROLLO",
        51, 27.7, 
        [
            "🔧 Necesidades: Desarrollo en competencias docentes y digitales",
            "📱 Tecnología: Brecha digital significativa (X6_CDD: -0.8)",
            "👨‍🏫 Capacitación: Docentes requieren fortalecimiento (-0.6)",
            "🎯 Oportunidad: Alto potencial con intervención adecuada",
            "📍 Contexto: Mixto rural-urbano con desafíos específicos"
        ],
        [
            "Red 47: 15 instituciones (29.4%)",
            "Red 79: 14 instituciones (27.5%)",
            "Red 48: 11 instituciones (21.6%)",
            "Red 44: 7 instituciones (13.7%)",
            "Red 72: 4 instituciones (7.8%)"
        ],
        [
            "👨‍🏫 Implementar fortalecimiento docente intensivo en competencias clave",
            "💻 Desarrollar programa específico contra la brecha digital",
            "📱 Mejorar infraestructura tecnológica y conectividad",
            "🎯 Focalizar en 3-4 competencias prioritarias por institución"
        ]
    )
    
    add_page_break(doc)
    
    # Tipología 5
    add_tipologia_section(
        doc, "5", "INSTITUCIONES URBANAS COMPLEJAS",
        11, 6.0,
        [
            "🏙️ Contexto: Urbano de alta densidad (X2_TR: -1.4 z-score)", 
            "👥 Complejidad: Ratio estudiante-docente alto (+2.1)",
            "🎯 Desafíos: Gestión de escala y atención diferenciada",
            "📊 Rendimiento: Variable según dimensión específica",
            "🔧 Especialización: Requerimientos urbanos específicos"
        ],
        [
            "Red 44: 6 instituciones (54.5%)",
            "Red 48: 3 instituciones (27.3%)",
            "Red 72: 2 instituciones (18.2%)"
        ],
        [
            "🏗️ Optimizar gestión organizacional para contextos de alta densidad",
            "👥 Desarrollar estrategias de atención diferenciada para ratios elevados", 
            "🎯 Implementar metodologías específicas para contexto urbano complejo",
            "🤝 Coordinar con Tipología 2 para transferir modelos urbanos exitosos"
        ]
    )
    
    # ANÁLISIS COMPARATIVO
    add_heading_with_color(doc, "VARIABLES MÁS DISCRIMINANTES", 1)
    
    doc.add_paragraph("Las siguientes variables muestran las mayores diferencias entre tipologías:")
    
    discriminantes = [
        "🌾 X2_TR (Ruralidad): Factor diferenciador principal entre contextos rurales y urbanos",
        "👥 X11_RED (Ratio Estudiante-Docente): Discrimina entre atención personalizada y masiva", 
        "👨‍🏫 X4_IDD (Desempeño Docente): Separa instituciones con fortaleza vs. necesidad docente",
        "💻 X6_CDD (Competencia Digital): Diferencia instituciones tecnológicamente avanzadas"
    ]
    
    for item in discriminantes:
        p = doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # DISTRIBUCIÓN TERRITORIAL
    add_heading_with_color(doc, "DISTRIBUCIÓN TERRITORIAL ESTRATÉGICA", 1)
    
    redes_info = [
        ("RED 44", "Urbano Alto Rendimiento", "57.9% Alto Rendimiento, 15.8% Urbanas Complejas"),
        ("RED 47", "Rural Diversificado", "45.0% Resilientes Rurales, 37.5% En Desarrollo"),
        ("RED 48", "Mixto Equilibrado", "47.1% Alto Rendimiento, 32.4% En Desarrollo"), 
        ("RED 54", "Equilibrado con Potencial", "53.3% Alto Rendimiento, 46.7% Equilibradas"),
        ("RED 72", "Alto Rendimiento Consolidado", "54.5% Alto Rendimiento, 18.2% Equilibradas"),
        ("RED 79", "Rural con Diversidad", "40.0% En Desarrollo, 34.3% Resilientes Rurales")
    ]
    
    for red, perfil, distribucion in redes_info:
        doc.add_heading(f"{red} - {perfil}", level=2)
        p = doc.add_paragraph(f"Distribución: {distribucion}")
    
    # CONCLUSIONES Y RECOMENDACIONES
    add_heading_with_color(doc, "CONCLUSIONES Y PRÓXIMOS PASOS", 1)
    
    doc.add_heading("Logros Conseguidos:", level=2)
    logros = [
        "✅ 6 tipologías institucionales robustas identificadas científicamente",
        "✅ Base metodológica excepcional con 21+ variables y validación cruzada",
        "✅ Estrategias específicas y aplicables para cada tipología",
        "✅ Red de 63 instituciones líderes para mentoría y transferencia"
    ]
    
    for logro in logros:
        p = doc.add_paragraph(logro, style='List Bullet')
    
    doc.add_heading("Recomendaciones Inmediatas:", level=2)
    recomendaciones = [
        "🔄 Validación con equipos territoriales Fe y Alegría (30 días)",
        "🎯 Diseño de intervenciones específicas por tipología (30 días)",
        "🚀 Pilotaje en instituciones representativas de cada tipo (30 días)",
        "📊 Implementación de sistema de monitoreo continuo (90 días)"
    ]
    
    for rec in recomendaciones:
        p = doc.add_paragraph(rec, style='List Bullet')
    
    # INFORMACIÓN TÉCNICA
    add_heading_with_color(doc, "INFORMACIÓN TÉCNICA", 1)
    
    tecnica = [
        "Base de datos: reasis_database_v3.db (184 instituciones)",
        "Variables: 21+ variables metodológicas y contextuales estandarizadas",
        "Metodología: K-Means + PCA + análisis correlacional + validación DBSCAN",
        "Calidad: Silhouette 0.194, Davies-Bouldin 1.623, Calinski-Harabasz 25.59"
    ]
    
    for item in tecnica:
        p = doc.add_paragraph(item, style='List Bullet')
    
    # Pie de página
    doc.add_paragraph()
    footer = doc.add_paragraph("PROYECTO REASIS - TIPOLOGÍAS INSTITUCIONALES FE Y ALEGRÍA 2025")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].italic = True
    
    subfooter = doc.add_paragraph("184 instituciones | 6 tipologías | 21+ variables | Metodología científica robusta")
    subfooter.alignment = WD_ALIGN_PARAGRAPH.CENTER  
    subfooter.runs[0].font.size = Pt(9)
    
    return doc

def main():
    """Función principal"""
    
    # Crear documento
    doc = create_executive_document()
    
    # Guardar en ubicación apropiada
    output_path = Path("assets/Consultoria/01 Informe en elaboración/INFORME_EJECUTIVO_TIPOLOGIAS_V3_FINAL.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_path))
    print(f"✅ Documento Word ejecutivo generado: {output_path}")
    
    # También guardar una copia en la carpeta de reportes
    backup_path = Path("data/reports/INFORME_EJECUTIVO_TIPOLOGIAS_V3_FINAL.docx")
    backup_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(backup_path))
    print(f"✅ Copia de respaldo generada: {backup_path}")

if __name__ == "__main__":
    main()