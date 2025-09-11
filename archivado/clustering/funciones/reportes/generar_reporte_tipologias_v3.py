#!/usr/bin/env python3
"""
Reporte Preliminar 3.0: aplica mejoras sobre una copia de la BD
- Copia segura: reasis_database_v3.db
- Correlación y PCA (reducción 90% varianza)
- Evaluación KMeans (K=3..6) con múltiples métricas y ranking
- Exploración DBSCAN
- Genera Word v3 y documento técnico Markdown
"""

from __future__ import annotations

import shutil
import sqlite3
from pathlib import Path
from typing import List, Tuple, Dict

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.cluster import KMeans, DBSCAN
from sklearn.decomposition import PCA
from sklearn.impute import SimpleImputer
from sklearn.metrics import silhouette_score, calinski_harabasz_score, davies_bouldin_score
from sklearn.preprocessing import StandardScaler

from docx import Document
from docx.shared import Inches


DB_ORIG = "reasis_database.db"
DB_V3 = "reasis_database_v3.db"
OUT_DIR = Path("data/reports/tipologias_v3")
FIG_DIR = OUT_DIR / "figs"
DOC_PATH = Path("assets/Consultoria/01 Informe en elaboración/REPORTE_PRELIMINAR_TIPOLOGIAS_2025_v3.docx")
DOC_TEC = Path("docs/INFORME_PRELIMINAR_3_0.md")


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    DOC_TEC.parent.mkdir(parents=True, exist_ok=True)


def copy_database():
    shutil.copyfile(DB_ORIG, DB_V3)


def columnas_para_modelo(conn: sqlite3.Connection) -> List[str]:
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(indices_metodologicos)")
    cols = {r[1] for r in cur.fetchall()}
    base = [
        "Y1_ILA","Y2_TD","Y3_PR","X1_NVC","X2_TR","X4_IDD","X6_CDD","X10_IE",
        "X11_RED","X11_RED_ajustado","X12_TOE","X13_TMATRC","X14_NIVEL_EDUCATIVO","X16_MODALIDAD",
        "X17_GESTION","X18_TURNO","X19_ORGANIZACION_PEDAGOGICA","X20_DIRECTIVOS_TOTAL","X21_MULTIPLICIDAD1",
        "X22_MULTIPLICIDAD2","X23_POBREZA_DISTRITO","X24_GPMD","X25_POBLACION_DISTRITO",
    ]
    if "X11_RED_ajustado" in cols and "X11_RED" in base:
        base.remove("X11_RED")
    return [c for c in base if c in cols]


def cargar_y_estandarizar(conn: sqlite3.Connection, cols: List[str]) -> Tuple[pd.DataFrame, pd.DataFrame, List[str]]:
    sel = ["CODIGO_MODULAR","NUMERO_FYA"] + cols
    df = pd.read_sql_query(f"SELECT {', '.join(sel)} FROM indices_metodologicos", conn)
    for c in cols:
        df[c] = pd.to_numeric(df[c], errors="coerce")
    usable = [c for c in cols if df[c].notna().sum() > 0 and df[c].dropna().nunique() > 1]
    imp = SimpleImputer(strategy="median")
    sc = StandardScaler()
    X_imp = pd.DataFrame(imp.fit_transform(df[usable]), columns=usable, index=df.index)
    X = pd.DataFrame(sc.fit_transform(X_imp), columns=usable, index=df.index)
    return df, X, usable


def plot_correlacion(df_num: pd.DataFrame, path: Path):
    corr = df_num.corr(numeric_only=True)
    plt.figure(figsize=(max(8, corr.shape[0]*0.3), max(6, corr.shape[1]*0.3)))
    sns.heatmap(corr, cmap="coolwarm", center=0)
    plt.title("Matriz de correlación (Pearson)")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    corr.to_csv(OUT_DIR/"correlacion_pearson.csv", index=True, encoding="utf-8")


def pca_90(X: pd.DataFrame) -> Tuple[np.ndarray, PCA, Path]:
    p = PCA()
    p.fit(X)
    csum = np.cumsum(p.explained_variance_ratio_)
    n = int(np.searchsorted(csum, 0.9) + 1)
    p = PCA(n_components=n).fit(X)
    Z = p.transform(X)
    plt.figure(figsize=(6,4))
    plt.plot(np.arange(1, len(csum)+1), csum, marker='o')
    plt.axhline(0.9, color='red', linestyle='--')
    plt.title(f"PCA - Varianza explicada (n={n})")
    plt.xlabel("Componentes")
    plt.ylabel("Varianza acumulada")
    plt.tight_layout()
    fig = FIG_DIR/"pca_varianza.png"
    plt.savefig(fig, dpi=180)
    plt.close()
    return Z, p, fig


def evaluar_kmeans(Z: np.ndarray) -> Tuple[pd.DataFrame, Path]:
    rows = []
    for k in range(3, 7):
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        lab = km.fit_predict(Z)
        sil = silhouette_score(Z, lab)
        db = davies_bouldin_score(Z, lab)
        ch = calinski_harabasz_score(Z, lab)
        rows.append({"K": k, "silhouette": sil, "davies_bouldin": db, "calinski_harabasz": ch})
    df = pd.DataFrame(rows)
    df["rank"] = df["silhouette"].rank(ascending=False) + df["calinski_harabasz"].rank(ascending=False) + df["davies_bouldin"].rank(ascending=True)
    df = df.sort_values("rank").reset_index(drop=True)
    df.to_csv(OUT_DIR/"kmeans_metricas_comparacion.csv", index=False, encoding="utf-8")
    plt.figure(figsize=(6,4))
    plt.plot(df['K'], df['silhouette'], marker='o', label='Silhouette')
    plt.plot(df['K'], (df['calinski_harabasz']-df['calinski_harabasz'].min())/(df['calinski_harabasz'].max()-df['calinski_harabasz'].min()+1e-9), label='CH (norm)')
    plt.plot(df['K'], 1-(df['davies_bouldin']-df['davies_bouldin'].min())/(df['davies_bouldin'].max()-df['davies_bouldin'].min()+1e-9), label='1-DB (norm)')
    plt.title('Comparativa de K (normalizada)')
    plt.xlabel('K')
    plt.legend()
    plt.tight_layout()
    fig = FIG_DIR/"kmeans_evaluacion.png"
    plt.savefig(fig, dpi=180)
    plt.close()
    return df, fig


def evaluar_dbscan(Z: np.ndarray) -> Tuple[Dict[str, float], Path]:
    std_avg = float(np.std(Z))
    eps = max(0.3, min(1.5, std_avg))
    dbs = DBSCAN(eps=eps, min_samples=5)
    lab = dbs.fit_predict(Z)
    n_noise = int((lab == -1).sum())
    n_clusters = int(len(set(lab)) - (1 if -1 in lab else 0))
    mask = (lab != -1)
    sil = float(silhouette_score(Z[mask], lab[mask])) if n_clusters >= 2 and mask.sum() > n_clusters else float("nan")
    resumen = {"eps": eps, "clusters": n_clusters, "ruido": n_noise, "silhouette": sil}
    pd.DataFrame([resumen]).to_csv(OUT_DIR/"dbscan_resumen.csv", index=False, encoding="utf-8")
    plt.figure(figsize=(5,3))
    plt.axis('off')
    plt.text(0.1, 0.6, f"DBSCAN\nclusters={n_clusters}\nruido={n_noise}\nsilhouette={sil:.3f}")
    fig = FIG_DIR/"dbscan_resumen.png"
    plt.savefig(fig, dpi=180)
    plt.close()
    return resumen, fig


def kmeans_final_guardar(conn: sqlite3.Connection, Z: np.ndarray, ids: pd.DataFrame, k: int) -> Tuple[int, float]:
    km = KMeans(n_clusters=k, random_state=42, n_init=10)
    lab = km.fit_predict(Z)
    sil = silhouette_score(Z, lab)
    cur = conn.cursor()
    cur.execute("DROP TABLE IF EXISTS resultados_clustering_v3")
    cur.execute("DROP TABLE IF EXISTS resumen_clusters_v3")
    cur.execute("CREATE TABLE resultados_clustering_v3 (codigo_modular TEXT PRIMARY KEY, numero_fya TEXT, cluster_asignado INTEGER, k_clusters INTEGER, silhouette_score REAL, fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
    cur.execute("CREATE TABLE resumen_clusters_v3 (cluster_id INTEGER PRIMARY KEY, n_instituciones INTEGER)")
    data = []
    for idx, l in zip(ids.index, lab):
        data.append((str(ids.at[idx, 'CODIGO_MODULAR']), str(ids.at[idx, 'NUMERO_FYA']), int(l), int(k), float(sil)))
    cur.executemany("INSERT OR REPLACE INTO resultados_clustering_v3 VALUES (?,?,?,?,?, CURRENT_TIMESTAMP)", data)
    uniq, cnt = np.unique(lab, return_counts=True)
    cur.executemany("INSERT OR REPLACE INTO resumen_clusters_v3 VALUES (?,?)", [(int(u), int(c)) for u,c in zip(uniq,cnt)])
    conn.commit()
    return int(k), float(sil)


def generar_word(resumen: Dict[str, str], figuras: Dict[str, Path]):
    doc = Document()
    doc.add_heading("REPORTE PRELIMINAR 3.0 - Tipologías Institucionales Fe y Alegría", level=1)
    p = doc.add_paragraph("Fecha: ")
    p.add_run(pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")).bold = True

    doc.add_heading("Resumen Ejecutivo", level=2)
    for k, v in resumen.items():
        para = doc.add_paragraph()
        para.add_run(f"{k}: ").bold = True
        para.add_run(v)

    doc.add_heading("Selección de Variables y PCA", level=2)
    if figuras.get("corr"):
        doc.add_paragraph("Correlación de variables (Pearson)")
        doc.add_picture(str(figuras["corr"]), width=Inches(5.5))
    if figuras.get("pca"):
        doc.add_paragraph("PCA - Varianza explicada")
        doc.add_picture(str(figuras["pca"]), width=Inches(5.5))

    doc.add_heading("Evaluación Multi-métrica y Alternativas", level=2)
    if figuras.get("k_eval"):
        doc.add_paragraph("Comparativa de K (Silhouette↑, CH↑, 1-DB↑)")
        doc.add_picture(str(figuras["k_eval"]), width=Inches(5.5))
    if figuras.get("dbscan"):
        doc.add_paragraph("Exploración DBSCAN")
        doc.add_picture(str(figuras["dbscan"]), width=Inches(5.5))

    if figuras.get("sizes"):
        doc.add_heading("Resultados Finales (v3)", level=2)
        doc.add_paragraph("Tamaños por cluster (v3)")
        doc.add_picture(str(figuras["sizes"]), width=Inches(5.5))

    doc.save(str(DOC_PATH))


def main():
    ensure_dirs()
    copy_database()

    conn = sqlite3.connect(DB_V3)
    try:
        cols = columnas_para_modelo(conn)
        df, X, usable = cargar_y_estandarizar(conn, cols)

        # Correlación
        corr_fig = FIG_DIR/"correlacion_pearson.png"
        plot_correlacion(df[usable], corr_fig)

        # PCA
        Z, pca, pca_fig = pca_90(X)

        # Evaluación de K y DBSCAN
        kdf, kfig = evaluar_kmeans(Z)
        dbs, dbfig = evaluar_dbscan(Z)

        # Selección simple por ranking
        k_sel = int(kdf.iloc[0]['K'])
        k_sel, sil_v3 = kmeans_final_guardar(conn, Z, df[['CODIGO_MODULAR','NUMERO_FYA']], k_sel)

        # Tamaños v3
        sizes = pd.read_sql_query("SELECT * FROM resumen_clusters_v3 ORDER BY cluster_id", conn)
        plt.figure(figsize=(6,4))
        sns.barplot(x='cluster_id', y='n_instituciones', data=sizes, palette='tab10')
        plt.title('Tamaños por cluster (v3)')
        plt.tight_layout()
        sizes_fig = FIG_DIR/"tamanos_cluster_v3.png"
        plt.savefig(sizes_fig, dpi=180)
        plt.close()

        resumen = {
            "Base de trabajo": "Copia reasis_database_v3.db (origen intacto)",
            "Variables usadas": f"{len(usable)}",
            "PCA componentes": f"{pca.n_components_} (~{int(np.round(np.sum(pca.explained_variance_ratio_)*100))}% varianza)",
            "K seleccionado": f"{k_sel}",
            "Silhouette (v3)": f"{sil_v3:.3f}",
            "DBSCAN (clusters/ruido)": f"{dbs['clusters']}/{dbs['ruido']}",
        }
        figuras = {"corr": corr_fig, "pca": pca_fig, "k_eval": kfig, "dbscan": dbfig, "sizes": sizes_fig}
        generar_word(resumen, figuras)

        # Doc técnico
        contenido = (
            f"# INFORME PRELIMINAR 3.0 - Detalle Técnico\n\n"
            f"Fecha: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}\n\n"
            f"Base: {DB_V3}\n\n"
            f"Variables candidatas: {len(cols)} | Usadas: {len(usable)}\n\n"
            f"PCA componentes: {pca.n_components_} | Varianza total: {np.sum(pca.explained_variance_ratio_):.3f}\n\n"
            f"KMeans evaluación (csv): data/reports/tipologias_v3/kmeans_metricas_comparacion.csv\n"
            f"DBSCAN resumen: {dbs}\n\n"
            f"K final: {k_sel} | Silhouette v3: {sil_v3:.3f}\n\n"
            f"Figuras en: data/reports/tipologias_v3/figs/\n"
        )
        DOC_TEC.write_text(contenido, encoding='utf-8')

        print(f"[OK] Reporte v3 generado: {DOC_PATH}")
        print(f"[OK] Informe técnico: {DOC_TEC}")
    finally:
        conn.close()


if __name__ == "__main__":
    main()


