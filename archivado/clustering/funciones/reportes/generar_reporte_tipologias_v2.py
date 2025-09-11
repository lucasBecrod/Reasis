#!/usr/bin/env python3
"""
Genera el REPORTE PRELIMINAR 2.0 (Word) de tipologías con K-Means avanzado.

Salidas:
- Word: assets/Consultoria/01 Informe en elaboración/REPORTE_PRELIMINAR_TIPOLOGIAS_2025_v2.docx
- Figuras PNG: data/reports/tipologias_v2/figs/
- CSVs resumen: data/reports/tipologias_v2/
"""

from __future__ import annotations

import os
import sqlite3
from pathlib import Path
from typing import List, Tuple, Dict

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.cluster import KMeans
from sklearn.impute import SimpleImputer
from sklearn.metrics import silhouette_score, calinski_harabasz_score, davies_bouldin_score, adjusted_rand_score
from sklearn.preprocessing import StandardScaler

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    Document = None


DB = "reasis_database.db"
OUT_DIR = Path("data/reports/tipologias_v2")
FIG_DIR = OUT_DIR / "figs"
DOC_PATH = Path("assets/Consultoria/01 Informe en elaboración/REPORTE_PRELIMINAR_TIPOLOGIAS_2025_v2.docx")

# Nombres legibles y notas pedagógicas
FRIENDLY = {
    "Y1_ILA": ("Logro académico (Y1)", "Síntesis del rendimiento estudiantil"),
    "Y2_TD": ("Tendencia de desempeño (Y2)", "Evolución del logro en el tiempo"),
    "Y3_PR": ("Progreso relativo (Y3)", "Desempeño esperado vs contexto"),
    "X1_NVC": ("Vulnerabilidad contextual (X1)", "Entorno socioeconómico y geográfico"),
    "X2_TR": ("Ruralidad (X2)", "Tipo de contexto territorial"),
    "X4_IDD": ("Desempeño docente (X4)", "Competencias y práctica pedagógica"),
    "X6_CDD": ("Competencia digital docente (X6)", "Uso educativo de tecnología"),
    "X10_IE": ("Infraestructura digital (X10)", "Conectividad y equipamiento"),
    "X11_RED": ("Ratio estudiantes/docente (X11)", "Carga docente"),
    "X11_RED_ajustado": ("Ratio ajustado (X11)", "Carga docente con tope pedagógico"),
    "X12_TOE": ("Organización escolar (X12)", "Unidocente/Multigrado/Polidocente"),
    "X13_TMATRC": ("Tendencia matrícula (X13)", "Evolución de matrícula"),
    "X14_NIVEL_EDUCATIVO": ("Nivel educativo (X14)", "Inicial/Primaria/Secundaria, etc."),
    "X16_MODALIDAD": ("Modalidad (X16)", "Escolarizada/No escolarizada"),
    "X17_GESTION": ("Gestión (X17)", "Pública directa/privada"),
    "X18_TURNO": ("Turno (X18)", "Horario de funcionamiento"),
    "X19_ORGANIZACION_PEDAGOGICA": ("Organización pedagógica (X19)", "Estructura docente"),
    "X20_DIRECTIVOS_TOTAL": ("Directivos (X20)", "Cantidad de directivos"),
    "X21_MULTIPLICIDAD1": ("Multiplicidad 1 (X21)", "Complejidad institucional"),
    "X22_MULTIPLICIDAD2": ("Multiplicidad 2 (X22)", "Complejidad institucional"),
    "X23_POBREZA_DISTRITO": ("Pobreza distrital (X23)", "Ranking de pobreza"),
    "X24_GPMD": ("Grupo pobreza monetaria (X24)", "Grupo territorial de pobreza"),
    "X25_POBLACION_DISTRITO": ("Población del distrito (X25)", "Tamaño poblacional"),
}


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)


def obtener_columnas_disponibles(conn: sqlite3.Connection) -> Tuple[List[str], List[str]]:
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(indices_metodologicos)")
    cols = [row[1] for row in cur.fetchall()]

    vars_met = [
        "Y1_ILA", "Y2_TD", "Y3_PR",
        "X1_NVC", "X2_TR", "X4_IDD", "X6_CDD", "X10_IE",
        "X11_RED", "X11_RED_ajustado", "X12_TOE", "X13_TMATRC",
    ]
    vars_ctx = [
        "X14_NIVEL_EDUCATIVO", "X16_MODALIDAD", "X17_GESTION", "X18_TURNO",
        "X19_ORGANIZACION_PEDAGOGICA", "X20_DIRECTIVOS_TOTAL", "X21_MULTIPLICIDAD1",
        "X22_MULTIPLICIDAD2", "X23_POBREZA_DISTRITO", "X24_GPMD", "X25_POBLACION_DISTRITO",
    ]
    met = [c for c in vars_met if c in cols]
    ctx = [c for c in vars_ctx if c in cols]
    if "X11_RED_ajustado" in met and "X11_RED" in met:
        met.remove("X11_RED")
    return met, ctx


def cargar_y_preparar(conn: sqlite3.Connection, columnas: List[str], ratio_min: float = 0.6) -> Tuple[pd.DataFrame, pd.DataFrame, List[str]]:
    sel = ["CODIGO_MODULAR", "NUMERO_FYA", "NOMBRE_INSTITUCION"] + columnas
    df = pd.read_sql_query(f"SELECT {', '.join(sel)} FROM indices_metodologicos", conn)

    # Coerción numérica
    for c in columnas:
        df[c] = pd.to_numeric(df[c], errors="coerce")

    # Filtrar columnas válidas (varianza > 0)
    cols_valid = [c for c in columnas if df[c].notna().sum() > 0 and df[c].dropna().nunique() > 1]
    df["_validas"] = df[cols_valid].notna().sum(axis=1)
    min_req = max(1, int(np.ceil(len(cols_valid) * ratio_min)))
    dff = df[df["_validas"] >= min_req].drop(columns=["_validas"]).copy()

    # Imputación y estandarización
    imputer = SimpleImputer(strategy="median")
    X_imp = pd.DataFrame(imputer.fit_transform(dff[cols_valid]), columns=cols_valid, index=dff.index)
    scaler = StandardScaler()
    X = pd.DataFrame(scaler.fit_transform(X_imp), columns=cols_valid, index=dff.index)
    return dff, X, cols_valid


def silhouette_por_k(X: pd.DataFrame, kmin: int = 2, kmax: int = 7) -> Dict[int, float]:
    out: Dict[int, float] = {}
    for k in range(kmin, min(kmax, len(X) - 1) + 1):
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        labels = km.fit_predict(X)
        try:
            out[k] = float(silhouette_score(X, labels))
        except Exception:
            out[k] = float("nan")
    return out


def robustez_multisemeilla(X: pd.DataFrame, base_labels: np.ndarray, k: int, n_seeds: int = 30) -> Tuple[pd.DataFrame, float, float]:
    registros = []
    for seed in range(n_seeds):
        km = KMeans(n_clusters=k, random_state=seed, n_init=10)
        labels = km.fit_predict(X)
        sil = silhouette_score(X, labels)
        ari = adjusted_rand_score(base_labels, labels)
        ch = calinski_harabasz_score(X, labels)
        db = davies_bouldin_score(X, labels)
        registros.append({"seed": seed, "silhouette": sil, "ARI": ari, "CH": ch, "DB": db})
    df = pd.DataFrame(registros)
    return df, float(df["silhouette"].mean()), float(df["ARI"].mean())


def plot_silhouette_vs_k(sil_por_k: Dict[int, float], path: Path):
    ks = sorted(sil_por_k.keys())
    vals = [sil_por_k[k] for k in ks]
    plt.figure(figsize=(6, 4))
    sns.lineplot(x=ks, y=vals, marker="o")
    plt.title("Silhouette por K")
    plt.xlabel("K")
    plt.ylabel("Silhouette")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()


def plot_tamanos_cluster(df_sizes: pd.DataFrame, path: Path):
    plt.figure(figsize=(6, 4))
    sns.barplot(x="cluster_id", y="n_instituciones", data=df_sizes, palette="tab10")
    plt.title("Tamaños por cluster")
    plt.xlabel("Cluster")
    plt.ylabel("Instituciones")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()


def plot_centroides_heatmap(df_centroides: pd.DataFrame, cols: List[str], path: Path):
    data = df_centroides.set_index("cluster_id")[cols]
    plt.figure(figsize=(max(6, len(cols) * 0.8), 4))
    sns.heatmap(data, annot=True, fmt=".2f", cmap="vlag", center=0)
    plt.title("Centroides (z-score) - variables clave")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()


def plot_redes_stack(conn: sqlite3.Connection, path: Path):
    q = """
    SELECT rc.cluster_asignado AS cluster, ie.NUMERO_FYA AS red, COUNT(*) as n
    FROM resultados_clustering_avanzado rc
    LEFT JOIN indices_metodologicos ie ON ie.CODIGO_MODULAR = rc.codigo_modular
    GROUP BY rc.cluster_asignado, ie.NUMERO_FYA
    """
    df = pd.read_sql_query(q, conn)
    tabla = df.pivot(index="red", columns="cluster", values="n").fillna(0).astype(int)
    tabla = tabla.sort_index()
    tabla.plot(kind="bar", stacked=True, figsize=(8, 4), colormap="tab20")
    plt.title("Distribución por red y cluster")
    plt.xlabel("Red FyA")
    plt.ylabel("Instituciones")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()


def _nivel_z(v: float) -> str:
    if v is None or pd.isna(v):
        return "Sin dato"
    if v >= 0.8:
        return "Muy Alto"
    if v >= 0.5:
        return "Alto"
    if v <= -0.8:
        return "Muy Bajo"
    if v <= -0.5:
        return "Bajo"
    return "Medio"


def _perfil_cluster_texto(fila: pd.Series) -> List[str]:
    bullets: List[str] = []
    # Mapear variables a etiquetas legibles
    mapping = [
        ("Y1_ILA", "Logro académico"),
        ("X1_NVC", "Vulnerabilidad contextual"),
        ("X2_TR", "Ruralidad"),
        ("X4_IDD", "Desempeño docente"),
        ("X10_IE", "Infraestructura digital"),
        ("X11_RED", "Ratio estudiante-docente"),
        ("X13_TMATRC", "Tendencia de matrícula"),
    ]
    for col, label in mapping:
        if col in fila.index:
            bullets.append(f"{label}: {_nivel_z(float(fila[col]))}")
    return bullets


def generar_perfiles_clusters(conn: sqlite3.Connection) -> List[Dict[str, object]]:
    df_cent = pd.read_sql_query("SELECT * FROM resumen_clusters_avanzado ORDER BY cluster_id", conn)
    perfiles: List[Dict[str, object]] = []
    for _, row in df_cent.iterrows():
        cluster_id = int(row["cluster_id"]) if "cluster_id" in row else int(row.name)
        n = int(row.get("n_instituciones", 0))
        bullets = _perfil_cluster_texto(row)
        perfiles.append({"cluster_id": cluster_id, "n": n, "bullets": bullets})
    return perfiles


def _variable_label(col: str) -> str:
    return FRIENDLY.get(col, (col, ""))[0]


def _narrativa_resumen(k_opt: int, silhouette_opt: float, df_sizes: pd.DataFrame, top_vars: List[str]) -> List[str]:
    textos: List[str] = []
    textos.append(
        f"Se identificaron {k_opt} clusters (tipologías) con una separación global (silhouette) de {silhouette_opt:.3f}. "
        "Este valor es coherente en contextos educativos multivariados, donde las dimensiones se solapan naturalmente."
    )
    sizes_txt = ", ".join([f"C{int(r.cluster_id)}: {int(r.n_instituciones)} IIEE" for _, r in df_sizes.iterrows()])
    textos.append(f"Distribución de tamaños por cluster: {sizes_txt}.")
    if top_vars:
        legibles = ", ".join([_variable_label(v) for v in top_vars])
        textos.append(f"Variables más discriminantes entre tipologías: {legibles} (ver análisis de relevancia).")
    textos.append(
        "Cada tipología representa un conjunto de condiciones compartidas útiles para planificar intervenciones: "
        "fortalecimiento docente, conectividad, reorganización operativa y estrategias de retención."
    )
    return textos


def analizar_relevancia_variables(X: pd.DataFrame, labels: np.ndarray, cols: List[str], df_base: pd.DataFrame) -> pd.DataFrame:
    # Usar X estandarizado para comparabilidad; eta^2 por variable
    df_res = []
    groups = pd.Series(labels, index=X.index)
    for col in cols:
        x = X[col].astype(float)
        mu = float(x.mean())
        ss_total = float(((x - mu) ** 2).sum())
        ss_between = 0.0
        for g, idx in groups.groupby(groups).groups.items():
            xg = x.loc[idx]
            if len(xg) == 0:
                continue
            d = float(xg.mean() - mu)
            ss_between += float(len(xg)) * (d ** 2)
        eta2 = ss_between / ss_total if ss_total > 1e-12 else 0.0

        # Estadísticos en escala original (df_base)
        if col in df_base.columns:
            s = pd.to_numeric(df_base[col], errors="coerce")
            med = float(s.median(skipna=True)) if s.notna().any() else np.nan
            p25 = float(s.quantile(0.25)) if s.notna().any() else np.nan
            p75 = float(s.quantile(0.75)) if s.notna().any() else np.nan
            std = float(s.std(skipna=True, ddof=0)) if s.notna().any() else np.nan
            try:
                moda_series = s.mode(dropna=True)
                moda = float(moda_series.iloc[0]) if not moda_series.empty else np.nan
            except Exception:
                moda = np.nan
        else:
            med = p25 = p75 = std = moda = np.nan

        nombre, _ = FRIENDLY.get(col, (col, ""))
        df_res.append({
            "variable": col,
            "nombre": nombre,
            "eta2": eta2,
            "mediana": med,
            "p25": p25,
            "p75": p75,
            "std": std,
            "moda": moda,
        })
    out = pd.DataFrame(df_res).sort_values("eta2", ascending=False).reset_index(drop=True)
    return out


def generar_word(
    resumen: Dict[str, str],
    figuras: Dict[str, Path],
    perfiles: List[Dict[str, object]],
    narrativa: List[str],
    topvars_df: pd.DataFrame,
    fig_vars_relevancia: Path | None,
):
    if Document is None:
        raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx")

    doc = Document()
    doc.add_heading("REPORTE PRELIMINAR 2.0 - Tipologías Institucionales Fe y Alegría", level=1)
    p = doc.add_paragraph("Fecha: ")
    p.add_run(pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")).bold = True

    doc.add_heading("Resumen Ejecutivo", level=2)
    for k, v in resumen.items():
        para = doc.add_paragraph()
        para.add_run(f"{k}: ").bold = True
        para.add_run(v)
    # Narrativa ampliada
    for t in narrativa:
        doc.add_paragraph(t)

    doc.add_heading("Resultados de Clustering", level=2)
    if figuras.get("silhouette_k"):
        doc.add_paragraph("Silhouette por K")
        doc.add_picture(str(figuras["silhouette_k"]), width=Inches(5.5))
    if figuras.get("sizes"):
        doc.add_paragraph("Tamaños por cluster")
        doc.add_picture(str(figuras["sizes"]), width=Inches(5.5))
    if figuras.get("centroides"):
        doc.add_paragraph("Centroides (z-score) - variables clave")
        doc.add_picture(str(figuras["centroides"]), width=Inches(5.5))
    if figuras.get("redes"):
        doc.add_paragraph("Distribución por red y cluster")
        doc.add_picture(str(figuras["redes"]), width=Inches(5.5))

    # Perfiles por cluster (lenguaje ejecutivo)
    doc.add_heading("Perfiles de Cluster (lenguaje ejecutivo)", level=2)
    doc.add_paragraph(
        "Cada cluster agrupa instituciones con patrones similares. Los niveles (Muy Bajo/Bajo/Medio/Alto/Muy Alto) "
        "se interpretan como desviaciones estándar relativas a la media de la red completa."
    )
    for perf in perfiles:
        doc.add_heading(f"Cluster {perf['cluster_id']} — {perf['n']} instituciones", level=3)
        for b in perf["bullets"]:
            doc.add_paragraph(b, style="List Bullet")
        # Sugerencias tácticas según señales frecuentes
        sugerencias: List[str] = []
        bullets = {b.split(':')[0]: b for b in perf["bullets"]}
        val = bullets.get("Ruralidad", "")
        if any(x in val for x in ["Alto", "Muy Alto"]):
            sugerencias.append("Adaptar estrategias rurales: multigrado, materiales contextualizados, itinerancias")
        val = bullets.get("Desempeño docente", "")
        if any(x in val for x in ["Bajo", "Muy Bajo"]):
            sugerencias.append("Plan de fortalecimiento docente focalizado (mentoría, acompañamiento en aula)")
        val = bullets.get("Infraestructura digital", "")
        if any(x in val for x in ["Bajo", "Muy Bajo"]):
            sugerencias.append("Mejoras de conectividad y equipamiento prioritario")
        val = bullets.get("Ratio estudiante-docente", "")
        if any(x in val for x in ["Alto", "Muy Alto"]):
            sugerencias.append("Optimizar distribución de docentes y reorganizar secciones")
        val = bullets.get("Tendencia de matrícula", "")
        if any(x in val for x in ["Bajo", "Muy Bajo"]):
            sugerencias.append("Acción de retención y captación: vínculo familiar-comunidad, pertinencia de oferta")
        if sugerencias:
            doc.add_paragraph("Sugerencias tácticas:")
            for s in sugerencias:
                doc.add_paragraph("- " + s)

    # Análisis de variables más relevantes
    doc.add_heading("Análisis de variables más relevantes", level=2)
    doc.add_paragraph(
        "Importancia estimada por variable (eta²) = proporción de varianza explicada por la separación en clusters. "
        "Valores cercanos a 1 indican fuerte discriminación entre tipologías."
    )
    if fig_vars_relevancia and fig_vars_relevancia.exists():
        doc.add_picture(str(fig_vars_relevancia), width=Inches(5.5))
    # Tabla top 10
    top10 = topvars_df.head(10)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "eta²"
    hdr[2].text = "Mediana"
    hdr[3].text = "P25–P75"
    hdr[4].text = "Desv. Est."
    hdr[5].text = "Moda"
    for _, r in top10.iterrows():
        row = table.add_row().cells
        row[0].text = r.get("nombre") or r.get("variable")
        row[1].text = f"{r['eta2']:.3f}"
        row[2].text = f"{r['mediana']:.2f}" if pd.notna(r['mediana']) else "-"
        row[3].text = (
            f"{r['p25']:.2f}–{r['p75']:.2f}" if pd.notna(r['p25']) and pd.notna(r['p75']) else "-"
        )
        row[4].text = f"{r['std']:.2f}" if pd.notna(r['std']) else "-"
        row[5].text = f"{r['moda']:.2f}" if pd.notna(r['moda']) else "-"

    doc.add_heading("Evaluación de Robustez", level=2)
    if figuras.get("robustez_silhouette"):
        doc.add_paragraph("Distribución de silhouette por semilla (K óptimo)")
        doc.add_picture(str(figuras["robustez_silhouette"]), width=Inches(5.5))
    for k in ["k_optimo", "silhouette_optimo", "silhouette_media_multi", "ARI_medio_multi"]:
        if k in resumen:
            para = doc.add_paragraph()
            para.add_run(f"{k.replace('_',' ').title()}: ").bold = True
            para.add_run(resumen[k])

    # Guía práctica: fundamentos y uso del resultado
    doc.add_heading("Fundamentos de K-Means (5 minutos)", level=2)
    doc.add_paragraph(
        "K-Means agrupa instituciones en K grupos minimizando la distancia dentro de cada grupo. "
        "Trabajamos con variables estandarizadas (z-score) para que cada indicador aporte en la misma escala."
    )
    doc.add_paragraph("Cómo leer los resultados:")
    items = [
        "Silhouette (0 a 1): separación entre clusters. ~0.18 indica separación razonable en contexto educativo real",
        "Centroides (heatmap): promedios estandarizados por cluster; valores >0.5 son 'Altos', < -0.5 'Bajos'",
        "Tamaños de cluster: priorizan dónde intervenir (clusters grandes = mayor cobertura)",
        "Distribución por red: dónde focalizar la gestión territorial",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("Cómo usar las tipologías en gestión", level=2)
    usos = [
        "Planificación por cluster: diseñar paquetes de intervención específicos por perfil",
        "Asignación de recursos: priorizar conectividad/equipamiento donde 'Infraestructura digital' sea Baja",
        "Gestión docente: mentoría y formación donde 'Desempeño docente' sea Bajo",
        "Equilibrio operativo: revisar carga y secciones donde 'Ratio estudiante-docente' sea Alto",
        "Sostenibilidad de matrícula: activar estrategias de retención donde 'Tendencia de matrícula' sea Baja",
    ]
    for u in usos:
        doc.add_paragraph(u, style="List Bullet")

    doc.add_heading("Buenas prácticas y limitaciones", level=2)
    doc.add_paragraph(
        "Buenas prácticas:")
    bps = [
        "Actualizar variables y re-ejecutar cada semestre para monitorear cambios",
        "Validar en territorio: contrastar perfiles con directores y equipos zonales",
        "No usar un único indicador para decidir: mirar el conjunto",
    ]
    for b in bps:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph("Limitaciones:")
    lims = [
        "Separación moderada (silhouette ~0.18) es esperable al mezclar múltiples dimensiones",
        "Clusters pequeños (p.ej., 4 IIEE) requieren validación cualitativa antes de escalar decisiones",
        "Resultados dependen de las variables disponibles y su calidad/actualidad",
    ]
    for l in lims:
        doc.add_paragraph(l, style="List Bullet")

    # Ejemplos de instituciones por cluster (z-scores clave)
    try:
        conn2 = sqlite3.connect(DB)
        df_examples = pd.read_sql_query(
            """
            SELECT cluster_asignado AS cluster,
                   codigo_modular,
                   NOMBRE_INSTITUCION,
                   Y1_ILA_ZS, X1_NVC_ZS, X2_TR_ZS, X4_IDD_ZS, X10_IE_ZS, X11_RED_ZS, X13_TMATRC_ZS
            FROM vw_clusters_zscores_avanzado
            ORDER BY cluster_asignado, ABS(Y1_ILA_ZS) DESC
            """,
            conn2,
        )
        conn2.close()
        doc.add_heading("Ejemplos por cluster (z-scores clave)", level=2)
        for cluster in sorted(df_examples['cluster'].unique()):
            sub = df_examples[df_examples['cluster'] == cluster].head(6)
            doc.add_heading(f"Cluster {int(cluster)} - 6 instituciones", level=3)
            table = doc.add_table(rows=1, cols=9)
            hdr = table.rows[0].cells
            hdr[0].text = "Código"
            hdr[1].text = "Institución"
            hdr[2].text = "Y1"
            hdr[3].text = "X1"
            hdr[4].text = "X2"
            hdr[5].text = "X4"
            hdr[6].text = "X10"
            hdr[7].text = "X11"
            hdr[8].text = "X13"
            for _, r in sub.iterrows():
                row = table.add_row().cells
                row[0].text = str(r['codigo_modular'])
                row[1].text = str(r['NOMBRE_INSTITUCION'])[:40]
                row[2].text = f"{r['Y1_ILA_ZS']:.2f}" if pd.notna(r['Y1_ILA_ZS']) else "-"
                row[3].text = f"{r['X1_NVC_ZS']:.2f}" if pd.notna(r['X1_NVC_ZS']) else "-"
                row[4].text = f"{r['X2_TR_ZS']:.2f}" if pd.notna(r['X2_TR_ZS']) else "-"
                row[5].text = f"{r['X4_IDD_ZS']:.2f}" if pd.notna(r['X4_IDD_ZS']) else "-"
                row[6].text = f"{r['X10_IE_ZS']:.2f}" if pd.notna(r['X10_IE_ZS']) else "-"
                row[7].text = f"{r['X11_RED_ZS']:.2f}" if pd.notna(r['X11_RED_ZS']) else "-"
                row[8].text = f"{r['X13_TMATRC_ZS']:.2f}" if pd.notna(r['X13_TMATRC_ZS']) else "-"
    except Exception:
        pass

    doc.save(str(DOC_PATH))


def main():
    ensure_dirs()
    conn = sqlite3.connect(DB)
    try:
        # Cargar columnas y matriz
        met, ctx = obtener_columnas_disponibles(conn)
        cols = met + ctx
        df_base, X, cols_used = cargar_y_preparar(conn, cols, ratio_min=0.6)

        # Silhouette por K
        sil_k = silhouette_por_k(X, 2, 7)
        fig_sil_k = FIG_DIR / "silhouette_por_k.png"
        plot_silhouette_vs_k(sil_k, fig_sil_k)

        # Cargar resultados base
        df_sizes = pd.read_sql_query("SELECT cluster_id, n_instituciones FROM resumen_clusters_avanzado ORDER BY cluster_id", conn)
        fig_sizes = FIG_DIR / "tamanos_cluster.png"
        plot_tamanos_cluster(df_sizes, fig_sizes)

        # Centroides (tomar desde tabla de resumen)
        df_cent = pd.read_sql_query("SELECT * FROM resumen_clusters_avanzado ORDER BY cluster_id", conn)
        key_cols = [c for c in ["Y1_ILA", "X1_NVC", "X2_TR", "X4_IDD", "X10_IE", "X11_RED", "X13_TMATRC"] if c in df_cent.columns]
        fig_cent = FIG_DIR / "centroides_heatmap.png"
        if key_cols:
            plot_centroides_heatmap(df_cent, key_cols, fig_cent)

        # Distribución por red
        fig_red = FIG_DIR / "distribucion_red_cluster.png"
        plot_redes_stack(conn, fig_red)

        # K óptimo detectado y labels base
        df_k = pd.read_sql_query("SELECT k_clusters, silhouette_score FROM resultados_clustering_avanzado LIMIT 1", conn)
        k_opt = int(df_k["k_clusters"].iloc[0]) if not df_k.empty else max(sil_k, key=sil_k.get)
        # Construir labels base a partir de resultados
        asign = pd.read_sql_query("SELECT codigo_modular, cluster_asignado FROM resultados_clustering_avanzado", conn)
        base_labels = asign.set_index("codigo_modular").loc[df_base["CODIGO_MODULAR"], "cluster_asignado"].values

        # Robustez multi-semilla
        df_rob, sil_mean, ari_mean = robustez_multisemeilla(X, base_labels, k_opt, n_seeds=30)
        df_rob.to_csv(OUT_DIR / "robustez_multisemilla.csv", index=False, encoding="utf-8")

        # Plot distribución silhouette
        plt.figure(figsize=(6, 4))
        sns.histplot(df_rob["silhouette"], bins=10, kde=True)
        plt.title(f"Silhouette (K={k_opt}) - 30 semillas")
        plt.xlabel("Silhouette")
        plt.ylabel("Frecuencia")
        plt.tight_layout()
        fig_rob = FIG_DIR / "robustez_silhouette.png"
        plt.savefig(fig_rob, dpi=180)
        plt.close()

        # Resumen para Word
        resumen = {
            "Instituciones analizadas": f"{len(df_base)}",
            "Variables usadas": f"{len(cols_used)}",
            "k_optimo": f"{k_opt}",
            "silhouette_optimo": f"{max(sil_k.values()):.3f}",
            "silhouette_media_multi": f"{sil_mean:.3f}",
            "ARI_medio_multi": f"{ari_mean:.3f}",
        }
        perfiles = generar_perfiles_clusters(conn)

        # Relevancia por variable
        relev = analizar_relevancia_variables(X, base_labels, cols_used, df_base)
        relev.to_csv(OUT_DIR / "variables_relevancia.csv", index=False, encoding="utf-8")
        # Plot top relevancia
        fig_vars = FIG_DIR / "variables_relevancia.png"
        try:
            topn = relev.head(12)
            plt.figure(figsize=(7, 4))
            sns.barplot(x="eta2", y="nombre", data=topn, orient="h", palette="viridis")
            plt.title("Variables más discriminantes (eta²)")
            plt.xlabel("Proporción de varianza explicada")
            plt.ylabel("")
            plt.tight_layout()
            plt.savefig(fig_vars, dpi=180)
            plt.close()
        except Exception:
            fig_vars = None

        # Narrativa del resumen
        top_vars_codes = relev.head(3)["variable"].tolist() if not relev.empty else []
        narrativa = _narrativa_resumen(k_opt, max(sil_k.values()), df_sizes, top_vars_codes)
        figuras = {
            "silhouette_k": fig_sil_k,
            "sizes": fig_sizes,
            "centroides": fig_cent if key_cols else None,
            "redes": fig_red,
            "robustez_silhouette": fig_rob,
        }

        generar_word(resumen, figuras, perfiles, narrativa, relev, fig_vars)
        print(f"[OK] Reporte generado: {DOC_PATH}")
    finally:
        conn.close()


if __name__ == "__main__":
    main()


